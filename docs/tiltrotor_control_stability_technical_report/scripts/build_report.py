#!/usr/bin/env python3
"""Build the editable Word control/stability technical report.

The script treats MATLAB CSV files as the numerical source of truth, retains
the validated model-principle chapters from the historical report, and builds
one editable Word deliverable.  It does not change any production model or
default parameter.
"""

from __future__ import annotations

import csv
import math
import re
import subprocess
import sys
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph


TITLE = "倾转旋翼机部件级飞行动力学建模、短舱动态状态扩展与开环操纵稳定特性分析研究报告"
ROOT = Path(__file__).resolve().parents[3]
REPORT = Path(__file__).resolve().parents[1]
HISTORICAL_MD = (
    ROOT
    / "docs"
    / "master_thesis_final_multiround"
    / "final"
    / "MASTER_THESIS_FINAL_CANDIDATE.md"
)
PANDOC = Path(r"C:\Program Files\Pandoc\pandoc.exe")
DOCX_OUT = REPORT / "TECHNICAL_REPORT_EDITABLE_REVISED.docx"
EXPECTED_DISPLAY_EQUATIONS = 89


REQUIRED_CSV = [
    "STATIC_STABILITY_DERIVATIVES.csv",
    "DAMPING_DERIVATIVES.csv",
    "DERIVATIVE_CROSSCHECK.csv",
    "CONTROL_EFFECTIVENESS_DERIVATIVES.csv",
    "CONTROL_DERIVATIVE_CROSSCHECK.csv",
    "MODAL_PARAMETERS.csv",
    "MODAL_PARTICIPATION.csv",
    "MODAL_CLASSIFICATION.csv",
    "MODAL_CONDITIONING.csv",
    "CONTROL_STEP_RESPONSE_METRICS.csv",
    "CONTROL_STEP_LINEAR_NONLINEAR_COMPARISON.csv",
    "CONTROL_STEP_TIMESTEP_CONVERGENCE.csv",
    "TRIM_CHARACTERISTICS_BY_MODE.csv",
    "REPRESENTATIVE_POINT_AUDIT.csv",
]


def rows(name: str) -> list[dict[str, str]]:
    path = REPORT / name
    with path.open("r", encoding="utf-8-sig", newline="") as stream:
        return list(csv.DictReader(stream))


def num(value: str | float | None) -> float:
    try:
        return float(value)  # type: ignore[arg-type]
    except (TypeError, ValueError):
        return math.nan


def yes(value: str | bool | None) -> bool:
    return str(value).strip().lower() in {"1", "true", "yes"}


def fmt(value: str | float | None, digits: int = 5) -> str:
    number = num(value)
    if math.isnan(number):
        text = "" if value is None else str(value)
        return text.replace("|", r"\|")
    if number == 0:
        return "0"
    if abs(number) >= 1e4 or abs(number) < 1e-3:
        return f"{number:.{digits}e}"
    return f"{number:.{digits}f}".rstrip("0").rstrip(".")


def tex_sci(value: float, digits: int = 3) -> str:
    mantissa, exponent = f"{value:.{digits}e}".split("e")
    return rf"{mantissa}\times10^{{{int(exponent)}}}"


def md_table(data: list[dict[str, str]], columns: list[tuple[str, str]], limit=None) -> str:
    if limit is not None:
        data = data[:limit]
    head = "| " + " | ".join(label for _, label in columns) + " |"
    rule = "|" + "|".join("---" for _ in columns) + "|"
    body = []
    for row in data:
        values = []
        for key, _ in columns:
            value = row.get(key, "")
            value = {
                "CREDIBLE": "可信",
                "FAILED": "不可行",
                "VALID_CENTRAL_DIFFERENCE": "有效中心差分",
                "true": "是",
                "false": "否",
                "1": "是" if key in {"nearRepeatedRoot", "pathologicalEigenvectors"} else "1",
                "0": "否" if key in {"nearRepeatedRoot", "pathologicalEigenvectors"} else "0",
                "NINE_STATE_PHYSICAL_CONTROL": "九状态直接物理操纵",
                "THIRTEEN_STATE_TORQUE": "十三状态短舱力矩输入",
                "THIRTEEN_STATE_ANGLE_COMMAND": "十三状态短舱角命令输入",
                "FULL_CROSSCHECK": "完整交叉核查",
                "MIXED_OR_UNCERTAIN_MODE": "混合或不确定模态",
                "helicopter_longitudinal": "直升机模式纵向配平",
                "conversion_longitudinal": "转换模式纵向配平",
                "airplane_longitudinal": "飞机模式纵向配平",
                "collective": "对称总距",
                "diffCollective": "差动总距",
                "cyclicLong": "对称纵向周期变距",
                "diffCyclic": "差动纵向周期变距",
                "aileron": "副翼",
                "elevator": "升降舵",
                "rudder": "方向舵",
                "CX_alpha": r"\(C_{X_\alpha}\)",
                "CZ_alpha": r"\(C_{Z_\alpha}\)",
                "Cm_alpha": r"\(C_{m_\alpha}\)",
                "CY_betaSlip": r"\(C_{Y_\beta}\)",
                "Cl_betaSlip": r"\(C_{l_\beta}\)",
                "Cn_betaSlip": r"\(C_{n_\beta}\)",
                "Cl_p": r"\(C_{l_p}\)",
                "Cl_r": r"\(C_{l_r}\)",
                "Cm_q": r"\(C_{m_q}\)",
                "Cn_p": r"\(C_{n_p}\)",
                "Cn_r": r"\(C_{n_r}\)",
                "Cm_elevator": r"\(C_{m_{\delta_e}}\)",
                "Cl_aileron": r"\(C_{l_{\delta_a}}\)",
                "Cn_rudder": r"\(C_{n_{\delta_r}}\)",
                "Cl_diffCollective": r"\(C_{l_{\Delta\theta_0}}\)",
                "Cn_diffCollective": r"\(C_{n_{\Delta\theta_0}}\)",
                "Cm_cyclicLong": r"\(C_{m_{\theta_{1s}}}\)",
                "Cl_diffCyclic": r"\(C_{l_{\Delta\theta_{1s}}}\)",
                "Cn_diffCyclic": r"\(C_{n_{\Delta\theta_{1s}}}\)",
            }.get(value, value)
            if key == "pointId":
                point = re.fullmatch(r"B(\d+)_V(\d+)", value or "")
                if point:
                    value = f"{int(point.group(1))}°/{int(point.group(2))} m/s"
            if re.fullmatch(r"[-+]?\d+(?:\.\d+)?(?:[eE][-+]?\d+)?", value or ""):
                value = fmt(value)
            values.append((value or "").replace("|", r"\|"))
        body.append("| " + " | ".join(values) + " |")
    return "\n".join([head, rule, *body])


def normalize_report_prose(markdown: str) -> str:
    replacements = {
        "`FULL_CROSSCHECK`": "完整交叉核查",
        "`MIXED_OR_UNCERTAIN_MODE`": "混合或不确定模态",
        "`physicalConverged=false`": "物理闭合检查未通过",
        "`coupledConverged`": "数值迭代收敛状态",
        "`physicalConverged`": "物理闭合状态",
        "`normalFlowRatioActual`": "局部法向流比",
        "`pitchCommand`": "虚拟俯仰命令",
        "`max(T,0)`": "\\(\\max(T,0)\\)",
        "`atLimit`": "达到控制边界",
        "`COLL`": "总距",
        "Berger文献": "Berger 文献",
        "本文采用V形验模思想": "本文采用 V 形验模思想",
        "模型内部统一使用SI制": "模型内部统一使用 SI 制",
        "任何NaN、Inf": "任何 NaN、Inf",
        "机体系采用右手系，x_b 轴向前、y_b 轴向右、z_b 轴向下；": "机体系采用右手系：机体纵轴 x_b 向前、机体横轴 y_b 向右、机体垂轴 z_b 向下；",
        "短舱角 β=0 表示旋翼轴接近垂直、处于直升机侧，β=π/2 表示旋翼轴接近机体纵轴、处于飞机侧。Berger 文献采用相反端点的角度记号，二者映射为 β=π/2-δ_nac，引用其结果时必须先执行该转换。": (
            "短舱角为 0 rad 时，旋翼轴接近垂直、处于直升机侧；短舱角为 90° 时，旋翼轴接近机体纵轴、处于飞机侧。Berger 文献采用相反端点的角度记号，本报告角度与该文献角度的换算为“本报告角度 = 90° − Berger 文献角度”，引用其结果时必须先执行该转换。"
        ),
        "地面系到机体系的姿态由3-2-1欧拉角描述。本文不在接近 θ=±π/2 的奇异区域形成结论。": "地面系到机体系的姿态由 3-2-1 欧拉角描述。本文不在俯仰角接近 ±90° 的奇异区域形成结论。",
        "这里 I 为关于实际整机重心、在机体系表达的完整对称惯量矩阵。": "上式中的惯量矩阵 I 为关于实际整机重心、在机体系表达的完整对称惯量矩阵。",
        "给出，其中 c_bullet=cos(bullet)、s_bullet=sin(bullet)，角度单位为 rad。": "给出，其中 c_•=cos(•)、s_•=sin(•)，角度单位为 rad。",
        "短舱角端点通过推力方向检查：β=0 时正推力主要沿 -z_b，\nβ=π/2 时主要沿 +x_b。重力端点通过水平姿态\ng_b=[0,0,g]^T 检查。": "短舱角端点通过推力方向检查：短舱角为 0 rad 时正推力主要沿机体垂轴负方向，短舱角为 90° 时正推力主要沿机体纵轴正方向。重力端点通过水平姿态下的机体系重力向量检查。",
        "长度m、质量kg、速度m/s、力N、力矩N·m、惯量kg·m²、角度rad、\n角速度rad/s": "长度 m、质量 kg、速度 m/s、力 N、力矩 N·m、惯量 kg·m²、角度 rad、角速度 rad/s",
        "以degree、rpm、英尺或磅力": "以 degree、rpm、英尺或磅力",
        "载荷合成。短舱角端点、左右旋向和SI单位": "载荷合成。短舱角端点、左右旋向和 SI 单位",
        "公开XV-15": "公开 XV-15",
        "采用atan2计算": "采用 atan2 计算",
        "但 物理闭合状态为 false": "但物理闭合状态为 false",
        "旋翼输出区分数值序列收敛 数值迭代收敛状态 与物理闭合 物理闭合状态": "旋翼计算分别检查数值迭代是否收敛和物理方程是否闭合",
        "h∈L,R": "h∈{L,R}",
        "s_h∈-1,+1": "s_h∈{-1,+1}",
        "盘面法向共同向 +e_D 倾斜": "盘面法向共同向 +e_D 方向倾斜",
        "C_l=C_lαα": "C_l=C_lα α",
        "自由流区和尾流区分别采用 V_∞,h 与\nV_slip,h": "自由流区和尾流区分别采用自由流速度 V_∞,h 与尾流区速度 V_slip,h",
        "χ=|V_n|/max(||V||,V_ref)": "χ=|V_n|/max(||V||, V_ref)",
        "这里的 χ 只表示机翼局部法向流比，对应代码中的 局部法向流比；转换配平的虚拟俯仰命令使用 虚拟俯仰命令": "局部法向流比 χ 用于机翼气动模型，虚拟俯仰命令用于转换模式配平；两者含义不同，计算过程也相互独立",
        "周期使两个盘面法向共同沿 +e_D 倾斜": "周期使两个盘面法向共同沿 +e_D 方向倾斜",
        "qSc_bar": "qS c_bar",
        "I^-1M": "I^-1 M",
        "μ_bar=(μ_L+μ_R)/2": "平均推进比 μ_bar=(μ_L+μ_R)/2",
        "因此使用 A(μ_L,β_L) 和\nA(μ_R,β_R) 分别计算": "因此分别使用覆盖函数 A(μ_L,β_L) 和 A(μ_R,β_R) 计算左右半翼覆盖面积",
        "若 A 为严格线性函数": "若覆盖函数 A 为严格线性函数",
        "其中 z 包括规定工况下的姿态和控制未知量，W_r 用于平衡平动加速度、角加速度和运动学约束的量级。": "其中，配平未知量 z 包括规定工况下的姿态和控制量；残差权重 W_r 用于平衡平动加速度、角加速度和运动学约束的量级。",
        "控制边界写为 z_min≤z≤z_max。": "控制边界写为 z_min≤z≤z_max。",
        "在可信配平点 (x_★,u_★) 附近，": "在可信配平点 x_star、u_star 附近，",
        "h_relmax": "h_rel max",
        "h_rel/2": "h_rel/2",
        "逆变换为 β_L=β_s+β_d、β_R=β_s-β_d。": "逆变换为左短舱角 β_L=β_s+β_d、右短舱角 β_R=β_s−β_d。",
        "当前 I_nac、ω_n、ζ、速率、加速度和转矩上限均为研究占位或": "当前短舱等效惯量 I_nac、自然频率 ω_n、阻尼比 ζ、速率、加速度和转矩上限均为研究占位或",
        "β_dot_h 有三类作用。": "短舱角速度 β_dot_h 有三类作用。",
        "β_dot_h,β_ddot_h 直接项置零后的载荷差。": "将 β_dot_h 与 β_ddot_h 直接项置零后的载荷差。",
        "A 和 B 分别是状态导数对状态和输入的 Jacobian": "状态矩阵 A 和输入矩阵 B 分别是状态导数对状态和输入的 Jacobian",
        "NASA TM-86854 的 NASA 字段 COLL 是": "NASA TM-86854 中的总距表示",
        "九状态控制顺序为 `collective`、`diffCollective`、`cyclicLong`、差动纵向周期变距、`aileron`、`elevator`、`rudder`。": "九状态控制依次为对称总距、差动总距、对称纵向周期变距、差动纵向周期变距、副翼、升降舵和方向舵。",
        "表中数值来自预先规定参数方案的完整非线性回代。它们只用于选择局部动态研究起点，不构成XV-15飞行配平数据。": (
            "表中数值来自预先规定参数方案的完整非线性回代。该表只用于说明局部动态研究起点，不构成 XV-15 飞行配平数据。"
        ),
        "这些定义不提供气动精度，却决定": "这些定义不提供气动精度，但决定",
        "两条路径共享几何和气动参数时的比较可以揭示数学实现差异，却不能恢复": "两条路径共享几何和气动参数时的比较可以揭示数学实现差异，但不能恢复",
        "峰值幅度却依赖": "峰值幅度依赖",
        "这两个条件能够排除非对称或非正定矩阵，却不能确认": "这两个条件能够排除非对称或非正定矩阵，但不能确认",
        "刚体，却不能预测": "刚体，但不能预测",
        "任何 NaN、Inf、复数": "任何非有限数值或复数",
        "这不是遗漏一个显然控制，而是当前左右旋翼稳态一阶挥舞映射的控制架构选择": "当前控制架构不包含独立横向周期输入，这是左右旋翼稳态一阶挥舞映射的既定选择",
        "它把对称纵向通道和反对称横航向通道显式分开": "该坐标变换把对称纵向通道和反对称横航向通道显式分开",
        "其执行机构相关特征根主要由概念带宽与阻尼参数决定": "十三状态模型的执行机构相关特征根主要由概念带宽与阻尼参数决定",
        "由于升降舵同时处于 达到控制边界 且": "由于升降舵已经达到控制边界，且",
        "Jacobian/SVD 条件": "雅可比矩阵与奇异值分解（SVD）条件",
        "该入口从默认 `params_berger13()` 参数集重新计算": "该入口调用默认参数函数 `params_berger13()`，重新计算",
        "报告构建入口为 `docs/tiltrotor_control_stability_technical_report/scripts/build_report.py`。": "Word 报告构建入口为 `docs/tiltrotor_control_stability_technical_report/scripts/build_report.py`。",
        "标记为 完整交叉核查": "标记为完整交叉核查",
        "标为 混合或不确定模态": "标为混合或不确定模态",
        "统一标记 混合或不确定模态": "统一标记为混合或不确定模态",
        "但 物理闭合标志": "但物理闭合标志",
        "文献中以 degree、rpm、英尺或磅力给出的量": "文献中以度、转每分钟、英尺或磅力给出的量",
        "采用 atan2 计算入流角": "采用双参数反正切函数计算入流角",
        "|B15/V20|": "|15°/20 m/s|",
        "|B45/V35|": "|45°/35 m/s|",
        "|B75/V80|": "|75°/80 m/s|",
        "B75/V040 与 B75/V060": "75°/40 m/s 与 75°/60 m/s",
        "本文工况远离该区域": "所选工况远离该区域",
        "本文工况避开该区域": "所选工况避开该区域",
        "以本文端点定义": "按上述端点定义",
        "两者含义不同，计算过程也相互独立，两者含义不同，计算过程也相互独立。": (
            "两者含义不同，计算过程也相互独立。"
        ),
        "本文 \\(\\chi_0=0.35\\)、\\(\\Delta\\chi=0.15\\) 均属概念模型设定": (
            "模型中的 \\(\\chi_0=0.35\\) 和 \\(\\Delta\\chi=0.15\\) 均为概念参数"
        ),
        "本文没有以绝对值、静默饱和或零替换来掩盖非物理结果。": (
            "计算中不使用绝对值、静默饱和或零值替换来掩盖非物理结果。"
        ),
        "单位为W。": "单位为 W。",
        "本文的左右旋向只影响有向面内力、反扭矩和谐波映射": (
            "左右旋向在当前模型中只影响有向面内力、反扭矩和谐波映射"
        ),
        "本文不把单一条件数阈值解释成普适物理标准；该条件数仅作为内部可信度判据。": (
            "单一条件数阈值不作为普适物理标准，仅用于判断局部数值敏感性。"
        ),
        "短舱质量位置随角度变化时，总重心和惯量必须与本次载荷调用使用的左右短舱角保持一致。若先用旧重心计算部件力臂、再更新质量属性，会人为引入一步延迟；在线性化中，这种顺序错误会表现为关于短舱角的伪导数。本文在每次状态函数入口先计算质量属性，再向所有部件传递统一重心。": (
            "短舱质量位置随角度变化时，总重心、惯量和部件载荷必须采用同一时刻的左右短舱角。"
            "若先用旧重心计算部件力臂，再更新质量属性，会人为引入一步延迟，并在线性化结果中形成短舱角伪导数。"
            "每个计算时刻均先更新质量属性，再以统一重心计算各部件载荷。"
        ),
        "本文通过带宽失配、速率失配和单侧延迟构造可控的不对称，但没有为真实左右差异赋值。": (
            "模型通过带宽失配、速率失配和单侧延迟描述可控的不对称，真实构型的左右差异尚未赋值。"
        ),
        "二阶执行机构没有型号辨识数据，所以本文不比较真实伺服品质，也不预测机械卡滞载荷。该限制在时域图和结论中保持一致。": (
            "二阶执行机构没有型号辨识数据，因此不用于比较真实伺服品质或预测机械卡滞载荷。"
        ),
        "这里的 \\(\\chi\\) 只表示机翼局部法向流比，对应代码中的 局部法向流比；转换配平的虚拟俯仰命令使用 虚拟俯仰命令": (
            "局部法向流比 \\(\\chi\\) 用于机翼气动模型；"
            "虚拟俯仰命令用于转换模式配平。两者含义不同，计算过程也相互独立"
        ),
        "九状态路径的共同短舱角适合完全对称准静态工况。它用": "九状态路径的共同短舱角适用于完全对称准静态工况，并采用",
        "模型没有独立横向周期输入。\n当前控制架构不包含独立横向周期输入，这是左右旋翼稳态一阶挥舞映射的既定选择": (
            "模型没有独立横向周期输入；该设置源于左右旋翼稳态一阶挥舞映射的既定控制架构"
        ),
        "评估完整载荷与把": "评估完整载荷与将",
        "分别是状态导数对状态和输入的 Jacobian": "分别为状态导数关于状态和输入的雅可比矩阵",
        "因此既有 MAE/RMSE 只能称为": "因此，既有平均绝对误差和均方根误差（MAE/RMSE）只能称为",
        "B 矩阵的 B9、十三状态短舱力矩输入和短舱角命令输入三列": "B 矩阵中的九状态输入矩阵（B9）、十三状态短舱力矩输入矩阵和短舱角命令输入矩阵",
        "B 路径按 B9、十三状态短舱力矩输入和十三状态短舱角命令输入三列分别核查": (
            "B 路径分别核查九状态输入矩阵（B9）、十三状态短舱力矩输入矩阵和十三状态短舱角命令输入矩阵"
        ),
        "现有B矩阵": "现有 B 矩阵",
        "rpm/rad·s⁻¹": "转每分钟/弧度每秒",
        "这里 \\(\\mathbf I\\) 为关于实际整机重心、在机体系表达的完整对称惯量矩阵。": (
            "惯量矩阵 \\(\\mathbf I\\) 是关于实际整机重心、在机体系表达的完整对称矩阵。"
        ),
        "Berger 文献采用相反端点的角度记号，二者映射为 \\(\\beta=\\pi/2-\\delta_{\\rm nac}\\)": (
            "Berger 文献采用相反端点的角度记号，本报告与该文献的换算关系为 \\(\\beta=\\pi/2-\\delta_{\\rm nac}\\)"
        ),
        "，而把它作为内部可信度守门量": "；该条件数仅作为内部可信度判据",
        "它们不是等概率样本": "三个代表点不是等概率样本",
        "它们决定重心移动量和惯量变化幅值": "这些概念参数决定重心移动量和惯量变化幅值",
        "它只在对称参数和小扰动附近成立": "镜像关系只在对称参数和小扰动附近成立",
        "它能回答速率、延迟和左右失配如何进入刚体": "该执行机构模型可用于分析速率、延迟和左右失配如何进入刚体运动",
        "本报告不把它扩展为负推力、风车或机械卡滞载荷模型": "本报告不将十三状态模型扩展为负推力、风车或机械卡滞载荷模型",
        "二者没有共享变量、传参或覆盖关系": "两者含义不同，计算过程也相互独立",
        "这里 $\\mathbf{I}$ 为": "惯量矩阵 $\\mathbf{I}$ 是",
        "当前移动短舱质量900 kg和质心半径0.75 m属于通用概念参数；": (
            "当前移动短舱质量为 900 kg，质心半径为 0.75 m，二者均属于通用概念参数；"
        ),
        "；它们决定": "；这些概念参数决定",
        "；它能回答": "；该执行机构模型可用于分析",
        "对称/差动坐标": "对称和差动坐标",
        "历史接口名 `diffCyclic` 在物理含义上指差动纵向周期变距。": "差动纵向周期变距表示左右旋翼的纵向周期输入等量反向变化。",
        "## 3.7 模型假设、耦合遗漏与适用范围": "## 3.7 模型假设与未建模效应",
        "## 4.7 特征根、模态与证据限制": "## 4.7 特征根与模态解释",
        "## 5.9 状态扩展的研究增量与单向边界": "## 5.9 状态扩展的作用与耦合范围",
        "本文采用 V 形验模思想：左侧把科学问题分解为坐标、力学、气动、质量和执行机构要求；底部由单元、守恒和数值收敛检查支撑；右侧依次执行独立代码交叉比较、文献基准和外部试验数据关联。外部数据无法形成同构输入时，证据不向上越级。": (
            "校核流程按 V 形结构组织。模型要求分解为坐标、力学、气动、质量和执行机构五部分，"
            "并通过单元检查、守恒关系和数值收敛检查逐项核对。在此基础上，再进行独立实现比较、"
            "文献趋势比较和外部数据对照。外部数据若与模型输入不对应，只比较变化方向，不用于评价整机预测误差。"
        ),
        "模型不允许用数值截断掩盖物理分支缺失。任何非有限数值或复数、迭代不收敛、越界插值或负推力未闭合都必须作为显式状态进入证据矩阵。": (
            "数值截断不能代替缺失的物理分支。非有限数值、复数、迭代不收敛、"
            "越界插值和负推力未闭合均作为单独的计算状态记录。"
        ),
        "论文因此分别使用“定义校核”“趋势”“限定工况定量结果”和“外部关联”表述。": (
            "下文根据比较目的，分别采用公式与接口核对、趋势比较、指定工况定量分析和外部数据对照。"
        ),
        "当前正式网格是显式指定模式的离散可行性证据，不构成跨模式边界的连续控制路径。": (
            "当前网格用于判断各显式配平模式下的离散可行性，不用于连接跨模式的连续控制路径。"
        ),
        "本章把第二章的统一定义落实为旋翼、左右机翼、机身和尾翼的部件级载荷链，最终输出关于实际重心的整机合力与合矩。本章方程链属于理论与实现定义；有限值、镜像、端点、力矩臂和多步长检查分别在第六、七章集中评价，因此不在各小节重复证据等级说明。": (
            "本章把第二章的统一定义落实到旋翼、左右机翼、机身和尾翼的载荷计算，"
            "最终得到关于整机重心的合力与合力矩。有限值、镜像关系、端点、力矩臂和多步长检查"
            "集中列于第六、七章，各部件小节不再重复。"
        ),
        "\\(\\max(T,0)\\) 仅定义当前正推力 Eq.13 的诱导速度目标": (
            "\\(\\max(T,0)\\) 只用于定义式（13）在正推力条件下的诱导速度目标"
        ),
        "必须返回“不支持的物理分支”，不得作为成功点": (
            "计算状态记为“负推力物理分支未闭合”，不作为收敛结果"
        ),
        "新增三点的跨子块范数均低于 $3.1\\times10^{-11}$，可作为左右符号、旋向和坐标变换的内部一致性证据。": (
            "三个代表点的跨子块范数均低于 $3.1\\times10^{-11}$，说明左右符号、旋向和坐标变换保持内部一致。"
        ),
        "表中数值来自预先规定参数方案的完整非线性回代。该表只用于说明局部动态研究起点，不构成 XV-15 飞行配平数据。": (
            "表中数值来自预先规定参数方案的完整非线性回代，用于确定局部动态分析的起点，并非 XV-15 飞行配平数据。"
        ),
        "模型层级由部件公式、整机载荷、九状态刚体、十三状态短舱扩展、配平与线性化组成。每向上一层，均保留下一层的输入输出和适用性状态。该结构使失败可以定位到具体部件或数值环节，避免以整机残差掩盖局部不收敛。": (
            "模型由部件公式、整机载荷、九状态刚体、十三状态短舱扩展、配平和线性化组成。"
            "各部分保留前一步的输入、输出和计算状态，便于把异常定位到具体部件或数值环节，"
            "避免整机残差掩盖局部不收敛。"
        ),
        "## 2.4 模型层级与信息流": "## 2.4 模型组织与信息流",
        "全文使用三个互有包含关系的动力学层级。准静态九状态模型把共同短舱角作为工况参数，适合配平、局部导数和既有结果回归。十三状态规定运动模型增加左右短舱角与角速度，适合研究带宽、速率和不同步历史。外部参考旋翼不是第四个整机层级，而是可以替换单一旋翼部件的对照实现。层级之间共享状态顺序、控制定义、部件接口和载荷合成点，差异项必须在模型入口处显式选择。": (
            "全文比较三种相关模型。准静态九状态模型把共同短舱角作为工况参数，用于配平、"
            "局部导数和既有结果回归。十三状态规定运动模型增加左右短舱角与角速度，用于研究带宽、"
            "速率和不同步过程。外部参考旋翼只替换单侧旋翼部件，不构成新的整机模型。"
            "三种模型采用相同的状态顺序、控制定义、部件接口和载荷合成点，差异项在模型入口处明确选择。"
        ),
        "这些定义不提供气动精度，但决定不同部件和模型层级能否无歧义比较。": (
            "这些定义不决定气动精度，但决定不同部件和模型之间能否直接比较。"
        ),
        "审计同时检查惯量对称性、特征值正定性和交叉惯量符号。": (
            "计算中同时检查惯量对称性、特征值正定性和交叉惯量符号。"
        ),
        "审计以人工改变作用点的方式验证力矩增量，能够发现叉乘顺序、参考点或坐标系错误。": (
            "通过人工改变作用点并核对力矩增量，可发现叉乘顺序、参考点或坐标系错误。"
        ),
        "报告因此分别使用“定义校核”“趋势”“限定工况定量结果”和“外部关联”表述。": (
            "后续分析按结果性质区分公式与接口核对、趋势比较、指定工况定量分析和外部数据对照。"
        ),
        "## 3.8 数值实现与失败保留": "## 3.8 数值实现与异常处理",
        "部件间主要通过局部速度、尾流覆盖、实际重心和总载荷耦合，没有全机CFD级干扰数据库。": (
            "部件间主要通过局部速度、尾流覆盖、实际重心和总载荷耦合，尚未引入高保真全机气动干扰数据库。"
        ),
        "负推力、诱导迭代失败、NaN、Inf和非预期复数均必须沿调用链返回。对正推力分支，配平可信度判据同时要求数值序列收敛与原始诱导闭合残差达标；对负推力分支，原载荷可以用于诊断，但在尚无风车/自转闭合时必须拒绝。配平求解器不能把这些点当作高惩罚但可接受的平衡，更不能以零载荷替代。保留失败点使可行区边界能够区分“控制触界”“数值失败”和“模型适用性不足”三类原因。": (
            "模型调用出现负推力、诱导迭代失败、非有限数值或非预期复数时，异常状态沿调用链返回。"
            "正推力工况只有在数值迭代收敛且原始诱导闭合残差达标后，才可进入配平结果。"
            "负推力载荷可以保留用于分析，但在风车或自转分支尚未闭合时不作为有效配平点，也不以零载荷替代。"
            "这样可以把可行区边界处的问题区分为控制触界、数值失败和模型适用范围不足。"
        ),
    }
    for old, new in replacements.items():
        markdown = markdown.replace(old, new)
    markdown = re.sub(r"(?<=[\u3400-\u9fff，。；：]) +(?=[\u3400-\u9fff])", "", markdown)
    markdown = re.sub(r"(?<=[\u3400-\u9fff，。；：])\n(?=[\u3400-\u9fff])", "", markdown)
    final_replacements = {
        "全文使用三个互有包含关系的动力学层级。准静态九状态模型把共同短舱角作为工况参数，适合配平、局部导数和既有结果回归。十三状态规定运动模型增加左右短舱角与角速度，适合研究带宽、速率和不同步历史。外部参考旋翼不是第四个整机层级，而是可以替换单一旋翼部件的对照实现。层级之间共享状态顺序、控制定义、部件接口和载荷合成点，差异项必须在模型入口处显式选择。": (
            "全文比较三种相关模型。准静态九状态模型把共同短舱角作为工况参数，用于配平、"
            "局部导数和既有结果回归。十三状态规定运动模型增加左右短舱角与角速度，用于研究带宽、"
            "速率和不同步过程。外部参考旋翼只替换单侧旋翼部件，不构成新的整机模型。"
            "三种模型采用相同的状态顺序、控制定义、部件接口和载荷合成点，差异项在模型入口处明确选择。"
        ),
        "这些定义不提供气动精度，但决定不同部件和模型层级能否无歧义比较。": (
            "这些定义不决定气动精度，但决定不同部件和模型之间能否直接比较。"
        ),
        "部件间主要通过局部速度、尾流覆盖、实际重心和总载荷耦合，没有全机CFD级干扰数据库。": (
            "部件间主要通过局部速度、尾流覆盖、实际重心和总载荷耦合，"
            "尚未引入高保真全机气动干扰数据库。"
        ),
        "负推力、诱导迭代失败、NaN、Inf和非预期复数均必须沿调用链返回。"
        "对正推力分支，配平可信度判据同时要求数值序列收敛与原始诱导闭合残差达标；"
        "对负推力分支，原载荷可以用于诊断，但在尚无风车/自转闭合时必须拒绝。"
        "配平求解器不能把这些点当作高惩罚但可接受的平衡，更不能以零载荷替代。"
        "保留失败点使可行区边界能够区分“控制触界”“数值失败”和“模型适用性不足”三类原因。": (
            "模型调用出现负推力、诱导迭代失败、非有限数值或非预期复数时，异常状态沿调用链返回。"
            "正推力工况只有在数值迭代收敛且原始诱导闭合残差达标后，才可进入配平结果。"
            "负推力载荷可以保留用于分析，但在风车或自转分支尚未闭合时不作为有效配平点，也不以零载荷替代。"
            "这样可以把可行区边界处的问题区分为控制触界、数值失败和模型适用范围不足。"
        ),
        "本文的左右旋向只影响有向面内力、反扭矩和谐波映射，不改变完全对称工况下单侧耗功的正值。": (
            "左右旋向在当前模型中只影响有向面内力、反扭矩和谐波映射，"
            "不改变完全对称工况下单侧耗功的正值。"
        ),
        "短舱质量位置随角度变化时，总重心和惯量必须与本次载荷调用使用的左右短舱角保持一致。"
        "若先用旧重心计算部件力臂、再更新质量属性，会人为引入一步延迟；在线性化中，"
        "这种顺序错误会表现为关于短舱角的伪导数。本文在每次状态函数入口先计算质量属性，"
        "再向所有部件传递统一重心。": (
            "短舱质量位置随角度变化时，总重心、惯量和部件载荷必须采用同一时刻的左右短舱角。"
            "若先用旧重心计算部件力臂，再更新质量属性，会人为引入一步延迟，并在线性化结果中形成短舱角伪导数。"
            "每个计算时刻均先更新质量属性，再以统一重心计算各部件载荷。"
        ),
        "本文通过带宽失配、速率失配和单侧延迟构造可控的不对称，但没有为真实左右差异赋值。": (
            "模型通过带宽失配、速率失配和单侧延迟描述可控的不对称，真实构型的左右差异尚未赋值。"
        ),
        "二阶执行机构没有型号辨识数据，所以本文不比较真实伺服品质，也不预测机械卡滞载荷。"
        "该限制在时域图和结论中保持一致。": (
            "二阶执行机构没有型号辨识数据，因此不用于比较真实伺服品质或预测机械卡滞载荷。"
        ),
    }
    for old, new in final_replacements.items():
        markdown = markdown.replace(old, new)
    return markdown


def normalize_word_math_source(source: str) -> str:
    """Normalize legacy TeX forms that Pandoc cannot convert to Word OMML."""
    source = re.sub(r"\{\\rm\s+([^{}]+)\}", r"\\mathrm{\1}", source)
    source = re.sub(r"\\rm\s+([A-Za-z]+)", r"\\mathrm{\1}", source)
    return source


def validate_report_language(source: str) -> None:
    marker = "\n# 第十三章"
    if marker not in source:
        raise RuntimeError("Appendix marker is missing from the Word report source.")
    body = source.rsplit(marker, 1)[0]
    forbidden_tokens = (
        "physicalConverged",
        "coupledConverged",
        "normalFlowRatioActual",
        "pitchCommand",
        "atLimit",
        "MIXED_OR_UNCERTAIN_MODE",
        "FULL_CROSSCHECK",
        "dampingRatio",
        "oscillationPeriodSeconds",
    )
    forbidden_phrases = (
        "证据层级",
        "本批次",
        "字段值",
        "触发病态判据",
        "NaN",
        "Inf",
        "CSV 保存",
        "全机CFD",
        "Eq.13",
    )
    found = [item for item in (*forbidden_tokens, *forbidden_phrases) if item in body]
    if "`" in body:
        found.append("正文代码反引号")
    if found:
        raise RuntimeError("Word report body contains internal or audit-style terms: " + ", ".join(found))


def extract_core_model_chapters() -> str:
    text = HISTORICAL_MD.read_text(encoding="utf-8")
    start = text.index("# 第二章")
    end = text.index("# 第六章")
    core = text[start:end]
    # The report is independent from the historical thesis and uses the new
    # task's figures.  Remove historical image/caption pairs but retain all
    # equations and model-principle prose.
    core = re.sub(r"\n!\[[^\]]*\]\([^)]+\)\n", "\n", core)
    core = re.sub(r"\n\*图\s*\d+[^*]*\*\n", "\n", core)
    historical_result_markers = (
        "联合优化",
        "三个新增工况",
        "三个点的九状态",
        "三个工况的该项",
        "本文代表点中",
        "B15和B45",
        "0.0079至0.2006",
        "75°构型俯仰平衡",
    )
    paragraphs = re.split(r"\n{2,}", core)
    core = "\n\n".join(
        paragraph
        for paragraph in paragraphs
        if not any(marker in paragraph for marker in historical_result_markers)
    )
    replacements = {
        "论文": "报告",
        "送审": "技术审查",
        "门禁": "可信度判据",
        "冻结": "预先规定",
        "PASS": "满足判据",
        "FAIL": "不满足判据",
        "production": "正式计算",
        "owner-visible": "可追溯",
    }
    for old, new in replacements.items():
        core = core.replace(old, new)
    trim_status_text = (
        "配平结果分为可信、边界敏感、不可行和数值失败四类。可信点满足残差、有限性、控制边界、"
        "雅可比条件和完整非线性回代要求。边界敏感点的残差仍可接受，但控制余度过低或条件数过大，"
        "只用于描述控制权限接近边界的情况。不可行点在物理限位内无法同时闭合目标方程。"
        "数值失败点存在求解器失败、部件迭代不收敛或非有限值。\n\n只有可信点进入定量动态分析。"
    )
    # Compatibility normalization for the historical source only.  The count
    # guard prevents a broad pattern from silently changing other chapters.
    core, status_replacements = re.subn(
        r"可信配平点必须同时满足：.*?本文使用.*?三\s*个中文类别。",
        trim_status_text,
        core,
        count=1,
        flags=re.DOTALL,
    )
    if status_replacements != 1:
        raise RuntimeError("Historical trim-status definition was not normalized exactly once.")

    trim_classification_text = (
        "本节沿用 4.1 节的四类状态，表 7.2 直接列出分类结果。分类时先检查求解器、部件迭代和有限性；"
        "存在异常时归为“数值失败”。随后检查完整非线性回代和目标方程；残差不达标或物理限位内"
        "无法闭合时归为“不可行”。残差可接受但控制余度过低或雅可比条件较差时归为“边界敏感”，"
        "其余点归为“可信”。\n\n只有可信点可作为定量时域和模态分析的初始状态。"
    )
    core, classification_replacements = re.subn(
        r"本文把配平结果分为.*?部件迭代不收敛或函数非有限。",
        trim_classification_text,
        core,
        count=1,
        flags=re.DOTALL,
    )
    if classification_replacements != 1:
        raise RuntimeError("Historical trim-classification definition was not normalized exactly once.")
    core = core.replace(
        "三层比较定义如下：第一层为九状态准静态短舱模型，短舱角作为固定参数；第二层为十三状态模型的对称流形，左右短舱状态相同；第三层允许左右短舱独立运动，并通过差动指令激发非对称响应。三个层级使用同一预先规定参数方案和同一配平状态，不重新优化参数。",
        "模型比较包括三种情况。九状态准静态模型把短舱角作为固定工况参数。十三状态模型的对称状态"
        "取左右短舱角和角速度相同。左右独立状态则允许差动指令引起非对称响应。三种情况采用相同的"
        "参数方案和配平状态，不重新优化参数。第 9 章讨论的短舱力矩输入和短舱角命令输入，是左右独立"
        "状态下的两种输入形式；对称状态是其对称子空间。",
    )
    return core.strip()


def select_mid(data: list[dict[str, str]], names: list[str]) -> list[dict[str, str]]:
    return [
        row
        for row in data
        if row.get("stepLevel") == "2" and row.get("coefficientName") in names
    ]


def build_markdown() -> str:
    rep = rows("REPRESENTATIVE_POINT_AUDIT.csv")
    trim = rows("TRIM_CHARACTERISTICS_BY_MODE.csv")
    static = rows("STATIC_STABILITY_DERIVATIVES.csv")
    damping = rows("DAMPING_DERIVATIVES.csv")
    controls = rows("CONTROL_EFFECTIVENESS_DERIVATIVES.csv")
    deriv_cross = rows("DERIVATIVE_CROSSCHECK.csv")
    control_cross = rows("CONTROL_DERIVATIVE_CROSSCHECK.csv")
    modal = rows("MODAL_PARAMETERS.csv")
    classes = rows("MODAL_CLASSIFICATION.csv")
    conditioning = rows("MODAL_CONDITIONING.csv")
    steps = rows("CONTROL_STEP_RESPONSE_METRICS.csv")
    step_compare = rows("CONTROL_STEP_LINEAR_NONLINEAR_COMPARISON.csv")
    step_dt = rows("CONTROL_STEP_TIMESTEP_CONVERGENCE.csv")

    key_static_names = [
        "CX_alpha",
        "CZ_alpha",
        "Cm_alpha",
        "CY_betaSlip",
        "Cl_betaSlip",
        "Cn_betaSlip",
    ]
    key_damping_names = ["Cl_p", "Cl_r", "Cm_q", "Cn_p", "Cn_r"]
    key_control_names = [
        "Cm_elevator",
        "Cl_aileron",
        "Cn_rudder",
        "Cl_diffCollective",
        "Cn_diffCollective",
        "Cl_diffCyclic",
        "Cn_diffCyclic",
        "Cm_cyclicLong",
    ]
    static_key = select_mid(static, key_static_names)
    damping_key = select_mid(damping, key_damping_names)
    control_key = select_mid(controls, key_control_names)
    full_a = sum(row.get("status") == "FULL_CROSSCHECK" for row in deriv_cross)
    partial_a = sum(row.get("status") == "PARTIAL_CROSSCHECK" for row in deriv_cross)
    b9_full = sum(row.get("statusB9") == "FULL_CROSSCHECK" for row in control_cross)
    b13_torque_full = sum(row.get("statusB13Torque") == "FULL_CROSSCHECK" for row in control_cross)
    b13_command_full = sum(row.get("statusB13Command") == "FULL_CROSSCHECK" for row in control_cross)
    b9_partial = sum(row.get("statusB9") == "PARTIAL_CROSSCHECK" for row in control_cross)
    b13_torque_partial = sum(row.get("statusB13Torque") == "PARTIAL_CROSSCHECK" for row in control_cross)
    b13_command_partial = sum(row.get("statusB13Command") == "PARTIAL_CROSSCHECK" for row in control_cross)
    unstable = sum(row.get("stability") == "UNSTABLE" for row in modal)
    pathological = sum(yes(row.get("pathologicalEigenvectors")) for row in conditioning)
    uncertain = sum(row.get("classification") == "MIXED_OR_UNCERTAIN_MODE" for row in classes)
    modal_keys = {
        (row.get("pointId"), row.get("modelKind"), row.get("modeIndex"))
        for row in modal
    }
    unstable_keys = {
        (row.get("pointId"), row.get("modelKind"), row.get("modeIndex"))
        for row in modal
        if row.get("stability") == "UNSTABLE"
    }
    uncertain_keys = {
        (row.get("pointId"), row.get("modelKind"), row.get("modeIndex"))
        for row in classes
        if row.get("classification") == "MIXED_OR_UNCERTAIN_MODE"
    }
    modal_overlap = len(unstable_keys & uncertain_keys)
    modal_total = len(modal_keys)
    model_root_counts = {
        model_kind: sum(row.get("modelKind") == model_kind for row in modal)
        for model_kind in (
            "NINE_STATE_PHYSICAL_CONTROL",
            "THIRTEEN_STATE_ANGLE_COMMAND",
            "THIRTEEN_STATE_TORQUE",
        )
    }
    class_by_key = {
        (row.get("pointId"), row.get("modelKind"), row.get("modeIndex")): row
        for row in classes
    }

    def nine_state_unstable_root(point_id: str, classification: str) -> float:
        matches = [
            row
            for row in modal
            if row.get("pointId") == point_id
            and row.get("modelKind") == "NINE_STATE_PHYSICAL_CONTROL"
            and row.get("stability") == "UNSTABLE"
            and class_by_key[
                (row.get("pointId"), row.get("modelKind"), row.get("modeIndex"))
            ].get("classification")
            == classification
        ]
        if len(matches) != 1:
            raise RuntimeError(
                f"Expected one {classification} unstable root at {point_id}, found {len(matches)}."
            )
        return num(matches[0].get("realPartPerSecond"))

    b15_lateral_root = nine_state_unstable_root(
        "B15_V020", "LATERAL_DIRECTIONAL_APERIODIC_MODE"
    )
    b15_longitudinal_root = nine_state_unstable_root(
        "B15_V020", "LONGITUDINAL_APERIODIC_MODE"
    )
    b45_longitudinal_root = nine_state_unstable_root(
        "B45_V035", "LONGITUDINAL_APERIODIC_MODE"
    )
    max_bio = max(num(row.get("biorthogonalityError")) for row in conditioning)
    max_dt = max(num(row.get("relativeChangeFromFinest")) for row in step_dt)
    max_dt_tex = tex_sci(max_dt)
    direction_fraction = sum(yes(row.get("directionAgreement")) for row in step_compare) / max(
        len(step_compare), 1
    )
    all_rep = all(yes(row.get("credible")) and yes(row.get("physicalConverged")) for row in rep)

    abstract = rf"""# 摘要与符号说明

本报告研究短舱角和空速变化对倾转旋翼机开环稳定性、操纵功效及短舱状态耦合的影响。分析对象为通用低阶部件级模型，包括九状态刚体模型和含左右短舱角、角速度的十三状态模型。代表工况为 15°/20 m/s、45°/35 m/s 和 75°/80 m/s，三点的配平可信度与旋翼物理闭合检查均“{'满足' if all_rep else '存在未满足项'}”。

计算采用保持空速大小不变的迎角和侧滑角扰动、三档中心差分、完整质量与惯量换算、特征根与参与因子分析，以及线性—非线性同时间步小扰动对照。

九状态结果给出本研究最直接的刚体稳定性结论：15°/20 m/s 工况存在实部为 {b15_lateral_root:.6f} s⁻¹ 的横航向非周期发散和实部为 {b15_longitudinal_root:.6f} s⁻¹ 的纵向非周期发散；45°/35 m/s 工况存在实部为 {b45_longitudinal_root:.6f} s⁻¹ 的纵向非周期发散；75°/80 m/s 工况除航向运动学零根外，非运动学刚体根均位于左半平面。静稳定与阻尼导数用于辨识恢复和耗散机制，操纵导数与阶跃响应用于比较转换过程中的控制权限重分配。这些结果共同给出开环增稳通道、控制分配和局部线性模型适用时间尺度的设计依据。

十三状态角命令模型在三个代表点均增加两组稳定的对称和差动执行机构极点 \(-3.2\pm2.4i\) s⁻¹，没有引入新的不稳定执行机构模态。九状态模型中的三个刚体失稳机理在状态扩展后仍然存在。力矩输入模型因短舱角坐标缺少恢复关系而产生附加零根和病态特征向量，因此只读取其特征根分布，不用个别参与因子解释物理模态。控制阶跃的线性—非线性响应方向一致比例为 {direction_fraction:.1%}，局部线性模型可用于判断各通道的响应方向和初始变化；峰值大小仍需结合相应有效时段分析。

**关键词：** 倾转旋翼机；部件级建模；开环稳定性；操纵导数；参与因子；短舱动态状态

## 主要符号

|符号|含义|单位|
|---|---|---|
|\(u,v,w\)|机体系速度分量|m/s|
|\(p,q,r\)|机体系角速度|rad/s|
|\(\phi,\theta,\psi\)|滚转、俯仰、偏航欧拉角|rad|
|\(x_b,y_b,z_b\)|机体纵轴、横轴和垂轴|m|
|\(\alpha,\beta_{{\\rm slip}}\)|迎角、侧滑角|rad|
|\(\beta_M\)|代表工况的短舱构型角|rad|
|\(\beta_L,\beta_R\)|左、右短舱角|rad|
|\(\beta_s,\beta_d\)|短舱对称、差动角坐标|rad|
|\(\dot\beta_h,\ddot\beta_h\)|第 \(h\) 侧短舱角速度、角加速度|rad/s，rad/s²|
|\(\theta_0,\Delta\theta_0\)|对称总距、差动总距|rad|
|\(\theta_{{1s}},\Delta\theta_{{1s}}\)|对称、差动纵向周期变距|rad|
|\(\delta_a,\delta_e,\delta_r\)|副翼、升降舵、方向舵偏角|rad|
|\(\mathbf F_b,\mathbf M_b\)|机体系合力、关于重心的合力矩|N，N·m|
|\(m,\mathbf I\)|整机质量、关于重心的惯量矩阵|kg，kg·m²|
|\(q_\infty,S,b,\bar c\)|自由流动压、参考面积、翼展、平均气动弦|Pa，m²，m，m|
|\(\mu,\chi\)|旋翼推进比、机翼局部法向流比|1|
|\(\mathbf A,\mathbf B\)|状态矩阵、输入矩阵|按状态与输入定义|
|\(\lambda=\sigma+j\omega\)|特征根及其实部、虚部|s⁻¹|
|\(\omega_n,\zeta,T\)|自然频率、阻尼比、振荡周期|rad/s，1，s|
|\(V,W\)|右、左特征向量矩阵|1|
|\(P_{{ki}}\)|状态 \(k\) 对模态 \(i\) 的参与因子|1|

## 主要缩写

|缩写|含义|
|---|---|
|BEMT|叶素动量理论|
|CG|重心|
|SVD|奇异值分解|
|NASA|美国国家航空航天局|

# 第一章　研究背景与技术目标

## 1.1 研究对象

倾转旋翼机在短舱转换过程中同时改变旋翼推力方向、机翼局部来流、尾翼动压、质量属性和操纵权限。公开数据不足以唯一辨识具体型号的完整气动和执行机构参数，因此本研究采用通用、低阶、部件级正向机理模型，强调坐标、单位、力矩参考点、物理分支与数值闭合的一致性。

## 1.2 技术目标与问题分解

研究围绕四个问题展开。首先建立统一的坐标系和部件载荷合成关系。随后在可信配平点计算静稳定、阻尼与操纵导数。九状态和十三状态模型用于分析局部特征根、状态参与情况及短舱执行机构模态。最后通过小幅操纵阶跃比较线性与非线性响应，并评估各控制通道的相对功效。

## 1.3 操纵稳定特性研究的意义

倾转过程中，旋翼推力方向、气动载荷分配和操纵通道同时变化，不能把直升机状态与固定翼状态作简单插值。同一控制输入在不同构型下可能承担不同的稳定和机动任务。

开环特征根用于确定需要增稳的状态组合及其变化时间尺度。静稳定和阻尼导数分别反映恢复作用与动态耗散，操纵导数反映各执行通道的直接控制能力。短舱状态扩展用于判断执行机构动态是否引入新的不稳定模态。上述结果可为增稳回路选择、控制分配调度、执行机构模型选型以及后续试验验证安排提供依据。

## 1.4 研究范围

计算采用给定的概念参数，并以三个可信配平点为局部分析基准。模型假设及结果外推范围统一列于 12.3 节。
""".replace("\\\\", "\\")

    chapter7 = rf"""# 第六章　模型校核与结果解释基础

## 6.1 校核内容与作用

模型校核用于排除控制映射、坐标变换、配平残差和差分步长造成的伪趋势。校核内容包括控制接口、极限工况、左右对称性、质量与惯量、旋翼物理闭合、配平回代、差分步长以及线性—非线性局部对照。文献曲线只用于比较变化方向，不改变内部计算结果的使用范围。

## 6.2 控制接口核查

九状态输入依次为对称总距、差动总距、对称纵向周期变距、差动纵向周期变距、副翼、升降舵和方向舵，短舱角作为外部构型参数给定。十三状态模型另设左右短舱力矩或左右短舱角命令。两类输入分别对应自由角力矩模型和规定运动角命令模型，不能混用。

## 6.3 NASA 旋翼曲线的比较范围

NASA TM-86854 以 \(0.75R\) 处的桨距表示总距，当前模型则把总距命令叠加到既定扭转分布上，两者的横坐标不能直接对应。这里只比较“总距增加时推力增加”的变化方向，不把曲线幅值差异作为预测误差。低总距计算失败点仍保留在原始结果中。

# 第七章　分模式配平特性

## 7.1 三个代表点

{md_table(rep, [
    ("pointId", "工况"),
    ("mode", "显式模式"),
    ("betaMDeg", "短舱角/(°)"),
    ("speedMps", "速度/(m/s)"),
    ("thetaDeg", "俯仰角/(°)"),
    ("alphaDeg", "迎角/(°)"),
    ("collectiveDeg", "总距/(°)"),
    ("cyclicLongDeg", "纵向周期/(°)"),
    ("elevatorDeg", "升降舵/(°)"),
    ("dynamicResidualNorm", "动态残差"),
    ("conditionNumber", "条件数"),
    ("minimumControlMarginFraction", "最小余度"),
])}

三点均采用显式配平模式并检查完整残差、控制边界、Jacobian/SVD 条件与左右旋翼物理闭合。75°/40 m/s 等失败点只保留在分模式离散表中，不进入导数、模态和阶跃分析。

## 7.2 分模式离散特性

{md_table(trim, [
    ("pointId", "工况"),
    ("mode", "模式"),
    ("speedMps", "速度"),
    ("betaMDeg", "短舱角"),
    ("thetaDeg", "俯仰角"),
    ("alphaDeg", "迎角"),
    ("collectiveDeg", "总距"),
    ("cyclicLongDeg", "纵向周期"),
    ("elevatorDeg", "升降舵"),
    ("minimumControlMarginFraction", "控制余度"),
    ("dynamicResidualNorm", "残差"),
    ("status", "状态"),
])}

![分模式离散配平特性](figures/FIG04_TRIM_CHARACTERISTICS_BY_MODE.png)

图中各模式分别连线，仅用于同一显式配平定义内的离散变化展示。跨模式边界不计算连续梯度，也不将九点集合解释为连续转换走廊。

B75/V040 与 B75/V060 的最小控制余度分别为 \(-9.57255\times10^{-10}\) 与 \(-4.55575\times10^{-10}\)。这两个极小负值来自浮点舍入。由于升降舵已经达到控制边界，且完整动态残差未达标，两个工况均归为“不可行”。

# 第八章　静稳定性、阻尼导数和操纵导数

## 8.1 定义与归一化

机体系采用 \(x\) 向前、\(y\) 向右、\(z\) 向下。迎角和侧滑角分别为

$$
\\alpha=\\operatorname{{atan2}}(w,u),\\qquad
\\beta_{{\\rm slip}}=\\arcsin(v/V).
$$

扰动迎角或侧滑角时保持 \(V=\sqrt{{u^2+v^2+w^2}}\) 不变。力系数以 \(q_\\infty S\) 归一化，滚转和偏航矩系数以 \(q_\\infty Sb\) 归一化，俯仰矩系数以 \(q_\\infty S\\bar c\) 归一化。角速度导数采用 \(pb/(2V)\)、\(q\\bar c/(2V)\)、\(rb/(2V)\) 的无量纲角速度。

中心差分为

$$
f_\\xi\\approx\\frac{{f(\\xi+h)-f(\\xi-h)}}{{2h}}.
$$

角度和物理控制采用 \(10^{{-3}},10^{{-4}},10^{{-5}}\) rad；角速度采用 \(10^{{-2}},10^{{-3}},10^{{-4}}\) rad/s。每个扰动量均计算正、负两个方向，并保留三档步长的有量纲导数、无量纲导数、离散差异和有效性状态。

## 8.2 关键静稳定导数

{md_table(static_key, [
    ("pointId", "工况"),
    ("coefficientName", "导数"),
    ("step", "中档步长"),
    ("dimensionalDerivative", "有量纲导数"),
    ("coefficientDerivative", "无量纲导数"),
    ("relativeStepVariation", "三步长相对差异"),
    ("status", "有效性"),
])}

## 8.3 关键阻尼导数

{md_table(damping_key, [
    ("pointId", "工况"),
    ("coefficientName", "导数"),
    ("step", "中档步长"),
    ("dimensionalDerivative", "有量纲导数"),
    ("coefficientDerivative", "无量纲导数"),
    ("relativeStepVariation", "三步长相对差异"),
    ("status", "有效性"),
])}

## 8.4 关键操纵导数

{md_table(control_key, [
    ("pointId", "工况"),
    ("coefficientName", "导数"),
    ("step", "中档步长"),
    ("coefficientDerivative", "无量纲导数"),
    ("relativeStepVariation", "三步长相对差异"),
    ("status", "有效性"),
])}

![三个代表点的关键直接导数](figures/FIG01_KEY_DERIVATIVES.png)

副翼、差动总距和差动纵向周期变距分别报告。对称纵向周期变距主要通过共同盘面倾斜进入俯仰通道；差动通道同时可能产生滚转和偏航交叉效应，符号必须结合左右分配和旋向解释。

## 8.5 A/B 矩阵交叉核查

状态矩阵 \(\mathbf A\) 和输入矩阵 \(\mathbf B\) 分别表示状态导数对状态量和输入量的偏导数，不能直接当作气动力系数导数。由直接力、力矩导数反推状态空间矩阵时，计算中计入质量、完整惯量矩阵、速度角关系和角速度运动学项。

本次计算对 \(\mathbf A\) 矩阵核查 {full_a} 项。九状态模型、十三状态力矩输入模型和十三状态角命令模型的 \(\mathbf B\) 矩阵各核查 {b9_full}、{b13_torque_full} 和 {b13_command_full} 项，结果均与直接导数换算一致。该核查表明，第 8.2—8.4 节所述的恢复、阻尼和操纵作用已经按相同的量纲及状态顺序进入线性模型。

# 第九章　特征根、模态参数及参与分析

## 9.1 参数定义

对复根 \(\lambda=\\sigma+j\\omega\)，采用

$$
\\omega_n=\\sqrt{{\\sigma^2+\\omega^2}},\\qquad
\\zeta=-\\frac{{\\sigma}}{{\\omega_n}},\\qquad
T=\\frac{{2\\pi}}{{|\\omega|}}.
$$

稳定实根报告 \(\tau=-1/\\lambda\)，不稳定根报告倍增时间 \(\ln 2/\\Re(\\lambda)\)。左右特征向量按 \(W^H V=I\) 双正交归一化，参与因子定义为 \(P_{{ki}}=V_{{ki}}W_{{ik}}\)，同时保存原始复数分量和按模态归一化的幅值。

## 9.2 条件性

{md_table(conditioning, [
    ("pointId", "工况"),
    ("modelKind", "模型"),
    ("matrixOrder", "阶数"),
    ("eigenvectorConditionNumber", "特征向量条件数"),
    ("biorthogonalityError", "双正交误差"),
    ("relativeMinimumSeparation", "最小相对间隔"),
    ("nearRepeatedRoot", "近重根"),
    ("pathologicalEigenvectors", "病态"),
])}

模态参数表中有 15 个阻尼比为空值，对应自然频率为零的根。实根的阻尼比按定义取 1，振荡周期则只对复根有意义，因此振荡周期共有 57 个空值。两列空值数量不同是定义域不同所致，不是数值计算失败。

本次计算中有 {pathological} 个十三状态矩阵出现近重根，其中三个力矩输入矩阵同时存在严重病态。三个工况的特征向量条件数分别为 \(5.10\times10^{{18}}\)、\(1.08\times10^{{18}}\) 和 \(1.75\times10^{{19}}\)，双正交误差均为 1.41421。力矩输入模型以短舱力矩驱动角速度，却没有为短舱角提供恢复关系，因此每个工况都出现两个附加零根。重复的自由角坐标使特征向量矩阵数值秩亏。双正交误差 1.41421 由伪逆计算直接得到，并非人为设置的上限。该模型的特征根分布可用于判断方程结构，但个别特征向量和参与因子不用于模态归因。

短舱角命令模型的特征向量条件数为 212.67–343.91，双正交误差保持在 \(10^{-15}\) 量级。这里的近重根来自两组相同的对称和差动执行机构极点 \(-3.2\pm2.4i\) s⁻¹，特征向量矩阵本身仍可用。由于重根子空间内的基底并不唯一，执行机构模态按对称和差动子空间解释，不再为其中每个向量赋予独立物理名称。

## 9.3 三类模型结果的使用方法

三个工况共计算 {modal_total} 个特征根，总数由 \((9+13+13)\times3\) 直接得到。这些根来自三类不同矩阵，物理用途并不相同。{model_root_counts['NINE_STATE_PHYSICAL_CONTROL']} 个九状态根用于分析刚体开环稳定性；{model_root_counts['THIRTEEN_STATE_ANGLE_COMMAND']} 个十三状态角命令根用于观察短舱规定运动状态加入后的耦合和执行机构子空间；{model_root_counts['THIRTEEN_STATE_TORQUE']} 个十三状态力矩输入根用于检查自由短舱角模型的方程结构，不用个别参与因子解释物理模态。

状态参与情况分为纵向刚体、横航向刚体、短舱对称、短舱差动和执行机构速率五组。主导组不明显、前两组接近、矩阵含近重根或特征向量条件较差时，将相应根归入混合或不确定模态。全部 {uncertain} 个十三状态根都进入这一类，原因是分类方法对矩阵级近重根采取了保守处理，并不表示每个根的特征值都失去意义。角命令模型按执行机构重根子空间解释，力矩输入模型只讨论特征根分布。其余根只有在状态参与足够集中时，才采用短周期、长周期、荷兰滚、螺旋或滚转收敛等名称。

## 9.4 正实部根及其主导状态

全部结果中共有 {unstable} 个正实部根，其中 {modal_overlap} 个同时归入混合或不确定模态。逐一比较工况和模型后可知，这 {unstable} 个根对应三个刚体非周期失稳机理，只是在三类矩阵中重复出现。

15°/20 m/s 九状态模型的横航向非周期根为 {b15_lateral_root:.7f} s⁻¹，主要状态为滚转角、偏航角速度和侧向速度。同一工况的纵向非周期根为 {b15_longitudinal_root:.7f} s⁻¹，主要状态为俯仰角、垂向速度和纵向速度。45°/35 m/s 的纵向非周期根为 {b45_longitudinal_root:.7f} s⁻¹，主要状态为俯仰角、纵向速度和垂向速度。75°/80 m/s 没有正实部的非运动学刚体根，航向零根属于运动学积分环节。

十三状态角命令模型保留了上述刚体根，新增的两组执行机构极点均位于左半平面，因此正实部根不是由短舱执行机构产生。15°/20 m/s 的增稳设计需要同时考虑横航向和纵向非周期发散，45°/35 m/s 主要考虑纵向非周期发散。75°/80 m/s 的局部结果没有显示正实部刚体根。力矩输入模型也给出相同位置的刚体根，但其特征向量病态，故不再用参与因子重复归因。

![九状态与十三状态开环特征根](figures/FIG02_OPEN_LOOP_EIGENVALUES.png)

![45°/35 m/s 十三状态命令模型参与因子](figures/FIG03_MODAL_PARTICIPATION.png)

不同显式配平模式之间未执行连续模态追踪。表中局部索引只在本工况、本模型内有效。

# 第十章　常规操纵响应

## 10.1 事件定义

对升降舵、副翼、方向舵、差动总距、差动纵向周期变距和对称纵向周期变距施加 0.5° 直接物理控制阶跃。非线性模型与九状态线性模型均采用二阶改进欧拉积分，并使用相同时间步。首先计算 0.02 s 和 0.01 s 两档时间步；当主响应峰值变化超过 2% 时，再增加 0.005 s 时间步。

## 10.2 峰值与相对效能

{md_table(steps, [
    ("pointId", "工况"),
    ("controlName", "输入"),
    ("stepAmplitudeDeg", "幅值/(°)"),
    ("dtSeconds", "最细步长/(s)"),
    ("primaryStateName", "主响应"),
    ("primaryPeak", "主峰值"),
    ("primaryPeakTimeSeconds", "峰值时间/(s)"),
    ("validDomainDurationSeconds", "有效域/(s)"),
    ("maximumLinearNonlinearStateError", "最大线性—非线性差"),
    ("relativeToB45V35", "相对45°/35"),
])}

![各通道直接物理控制相对效能](figures/FIG05_CONTROL_EFFECTIVENESS.png)

## 10.3 收敛性与结果解释

最细时间步的最大相对差异为 \({max_dt_tex}\)，线性与非线性状态峰值的方向一致率为 {direction_fraction:.1%}。初始角加速度反映状态方程的瞬时响应，有限时间内的角速度和姿态峰值还包含状态耦合。两类结果与静态力、力矩操纵导数含义不同。短舱角命令经过执行机构动态，其响应与本章的直接物理控制阶跃分别分析。

# 第十一章　短舱动态状态与开环响应

十三状态模型在九个刚体状态之外，引入左右短舱角和角速度。对称坐标描述共同倾转，差动坐标描述左右不同步。模型同时保留二阶规定运动执行机构、移动质量、反作用力矩和转子陀螺项。

三个代表点均出现两组稳定的执行机构极点 \(-3.2\pm2.4i\) s⁻¹，没有新增正实部执行机构模态。九状态模型中的刚体不稳定根在状态扩展后仍然存在。因此，当前概念执行机构既没有消除原有刚体发散，也不是这些发散的来源。

固定短舱角时，九状态与十三状态模型具有相同的静态载荷基准。短舱角速度扰动会重新计算动态载荷，并改变刚体子矩阵，因此十三状态扩展并不是简单增加四个执行机构极点。

力矩输入模型出现两个附加零根，双正交残差达到 \(\sqrt{2}\) 量级。这说明短舱角作为自由坐标时，需要补充角度恢复关系、约束方程或等价降阶，之后才能解释个别参与因子。现阶段的局部耦合研究采用角命令规定运动模型；自由角力矩输入模型需先完成方程闭合。

# 第十二章　讨论、结论与后续工作

## 12.1 主要结论

1. 三个代表点均满足配平可信度要求和左右旋翼物理闭合要求，可作为局部导数、模态和阶跃分析的基准。
2. 15°/20 m/s 同时存在横航向和纵向非周期发散，45°/35 m/s 存在纵向非周期发散。75°/80 m/s 的非运动学刚体根均位于左半平面。转换前段和中段的增稳设计应分别针对上述状态组合。
3. 静稳定和阻尼导数反映恢复与耗散作用，操纵导数和阶跃响应反映控制权限随构型的变化。两类结果可用于选择增稳反馈量，比较常规舵面与旋翼差动通道，并确定需要进行控制分配调度的工况。
4. 十三状态角命令模型只增加稳定的对称和差动执行机构极点，原有刚体失稳仍然存在。力矩输入模型的病态来自自由短舱角缺少恢复关系。短舱状态和输入的定义会直接影响模态结果的物理解释。
5. 六类物理操纵阶跃给出了初始方向、峰值、有效时段和时间步收敛结果。线性与非线性响应的方向一致率为 {direction_fraction:.1%}。局部线性模型适合判断响应方向和初始变化，峰值大小需按通道和有效时段分别分析。

## 12.2 操稳特性结论

短舱角和空速变化会同时改变自然稳定性与可用控制权限。三个代表点的纵向迎角导数均呈恢复方向，俯仰角速度导数均呈阻尼方向，但 15°/20 m/s 和 45°/35 m/s 仍存在刚体非周期发散。单个导数描述局部载荷作用，完整特征根还包含状态间耦合，因此静态恢复性不能单独保证动态稳定。

差动总距和差动纵向周期变距的局部功效随短舱角向飞机侧变化而减弱，横航向控制不宜沿用单一固定增益或固定分配。开环根确定需要抑制的状态组合，操纵导数确定可用的直接力矩通道。十三状态对照则把执行机构极点与刚体失稳分开，为短舱执行机构模型选型提供依据。

## 12.3 适用范围与局限

模型采用概念参数、稳态低阶旋翼模型和简化干扰关系，三个代表点只描述相应配平点附近的小扰动特性。现有数据尚不足以建立具体型号的参数辨识结果和连续转换走廊，也未开展跨模式连续模态追踪、操纵品质等级评定或伺服—铰链—结构双向载荷验证。因此，结论用于比较既定方程下的失稳状态、控制权限变化和短舱模型形式，不用于具体型号的稳定性或飞行安全包线判断。

## 12.4 后续工作

后续工作包括三项。第一，在 15°/20 m/s 和 45°/35 m/s 附近增加同模式局部点，追踪非周期发散随短舱角和空速的变化。第二，研究差动旋翼通道与常规舵面的控制分配调度。第三，为力矩输入短舱模型补充角度恢复关系，或采用约束、描述器和子空间降阶方法。待获得来源明确的公开参数、同构控制坐标和独立飞行辨识数据后，再评价这些结论对具体构型的适用程度。

# 第十三章　附录与复现说明

## 13.1 状态与输入顺序

九状态顺序为 \([u,v,w,p,q,r,\\phi,\\theta,\\psi]^T\)。控制量依次为对称总距、差动总距、对称纵向周期变距、差动纵向周期变距、副翼、升降舵和方向舵。为便于复现，相应代码接口名称依次为 `collective`、`diffCollective`、`cyclicLong`、`diffCyclic`、`aileron`、`elevator` 和 `rudder`。十三状态模型另含左右短舱角、角速度及相应的力矩输入或角命令输入。

## 13.2 复现入口

数值计算由 MATLAB 批处理脚本 `analysis/control_stability/run_full_assessment_batch.m` 完成。该脚本调用默认参数函数 `params_berger13()`，重新计算九点显式模式数据库，并在三个代表点生成数值结果、原始时历和图件。Word 报告由 `docs/tiltrotor_control_stability_technical_report/scripts/build_report.py` 生成。

## 13.3 结果文件

原始数据包含静稳定导数、阻尼导数、操纵导数、矩阵交叉核查、模态参数、参与因子、模态分类、矩阵条件性、控制阶跃指标、线性与非线性对照、时间步收敛和分模式配平特性。相应文件保留在报告目录中。

## 13.4 详细记录索引

详细的公式实现对应关系、数值测试记录和 Word 构建检查见 `TEST_REPORT.md`、`NUMERIC_OUTPUT_AUDIT.md` 和 `REPORT_BUILD_AND_QA.md`，此处不再展开。

## 13.5 参考资料

- Felker, F. F., Young, L. A., and Signor, D. B. *Full-Scale Hover Testing of the XV-15 Advanced Technology Blade Rotor*. NASA TM-86854, 1987.
- NASA TM-X-62407 与 NASA TM-81244：公开倾转旋翼机几何、气动与试验资料候选；数值引用须逐项复核构型、页码和单位。
- 南航公开论文 `NUAA_main_paper.pdf`：部件划分、坐标变换、六自由度方程、配平与数值线性化的方法参考。
""".replace("\\\\", "\\")

    markdown = "\n\n".join([abstract, extract_core_model_chapters(), chapter7]).rstrip() + "\n"
    return normalize_report_prose(markdown)


def pandoc_latex_body(markdown_path: Path) -> str:
    raise RuntimeError("Markdown/LaTeX report output is disabled; build the Word report only.")
    command = [
        str(PANDOC),
        str(markdown_path),
        "--from=markdown+tex_math_dollars",
        "--to=latex",
        "--top-level-division=chapter",
        f"--resource-path={REPORT}",
    ]
    result = subprocess.run(command, text=True, encoding="utf-8", capture_output=True)
    if result.returncode:
        raise RuntimeError(result.stderr)
    return result.stdout


def build_tex(markdown_path: Path) -> None:
    raise RuntimeError("Markdown/LaTeX report output is disabled; build the Word report only.")
    TEX_DIR.mkdir(parents=True, exist_ok=True)
    body = pandoc_latex_body(markdown_path)
    body = body.replace(r"\tightlist", "")
    preamble = rf"""\documentclass[UTF8,11pt,oneside,openany,fontset=windows]{{ctexbook}}
\usepackage[a4paper,left=24mm,right=24mm,top=25mm,bottom=24mm,headheight=15pt]{{geometry}}
\usepackage{{fontspec,xeCJK,graphicx,booktabs,longtable,array,amsmath,amssymb,bm}}
\usepackage{{caption,float,fancyhdr,hyperref,bookmark,microtype,xcolor}}
\graphicspath{{{{../}}}}
\usepackage[backend=biber,style=gb7714-2015,sorting=none]{{biblatex}}
\addbibresource{{references.bib}}
\setmainfont{{Times New Roman}}
\hypersetup{{hidelinks,pdftitle={{{TITLE}}},pdfauthor={{技术研究报告}}}}
\captionsetup{{font=small,labelsep=quad}}
\setlength{{\parindent}}{{2em}}
\setlength{{\parskip}}{{0pt}}
\linespread{{1.32}}
\raggedbottom
\pagestyle{{fancy}}
\fancyhf{{}}
\fancyhead[C]{{\small 倾转旋翼机开环操纵稳定特性分析研究报告}}
\fancyfoot[C]{{\thepage}}
\setcounter{{secnumdepth}}{{3}}
\setcounter{{tocdepth}}{{2}}
\emergencystretch=3em
\begin{{document}}
\begin{{titlepage}}
\centering
\vspace*{{28mm}}
{{\zihao{{1}}\bfseries\color[HTML]{{2E74B5}} {TITLE}\par}}
\vspace{{14mm}}
\rule{{0.78\textwidth}}{{0.8pt}}\par
\vspace{{12mm}}
{{\zihao{{3}} 通用低阶部件级模型技术研究报告\par}}
\vfill
{{\zihao{{-4}} 数值基线：MATLAB R2021a\par}}
\vspace{{4mm}}
{{\zihao{{-4}} 生成日期：{datetime.now():%Y年%m月%d日}\par}}
\end{{titlepage}}
\frontmatter
\tableofcontents
\listoffigures
\listoftables
\mainmatter
"""
    ending = r"""
\backmatter
\nocite{*}
\printbibliography[heading=bibintoc,title={参考文献}]
\end{document}
"""
    TEX_OUT.write_text(preamble + body + ending, encoding="utf-8")
    BIB_OUT.write_text(
        """@techreport{felker1987,
  author={Felker, Fort F. and Young, Larry A. and Signor, David B.},
  title={Full-Scale Hover Testing of the XV-15 Advanced Technology Blade Rotor},
  institution={NASA},
  number={TM-86854},
  year={1987}
}
@techreport{nasa62407,
  author={{NASA}},
  title={Tiltrotor Research Reference, NASA TM-X-62407},
  institution={NASA},
  number={TM-X-62407},
  year={1970}
}
@techreport{nasa81244,
  author={{NASA}},
  title={Tiltrotor Research Reference, NASA TM-81244},
  institution={NASA},
  number={TM-81244},
  year={1980}
}
""",
        encoding="utf-8",
    )


def set_style_font(style, east_asia: str, latin: str, size: float, bold=None, color=None):
    style.font.name = latin
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr, value in (("eastAsia", east_asia), ("ascii", latin), ("hAnsi", latin)):
        rfonts.set(qn(f"w:{attr}"), value)


def set_run_font(run, east_asia="微软雅黑", latin="Calibri", size=11, bold=None, color=None):
    run.font.name = latin
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr, value in (("eastAsia", east_asia), ("ascii", latin), ("hAnsi", latin)):
        rfonts.set(qn(f"w:{attr}"), value)


def insert_before(anchor: Paragraph, text="", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addprevious(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def add_field(paragraph: Paragraph, instruction: str, placeholder: str = ""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def set_table_geometry(table):
    widths_total = 9360
    count = len(table.columns)
    widths = [widths_total // count] * count
    widths[-1] += widths_total - sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(widths_total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[col_index]))
            tc_w.set(qn("w:type"), "dxa")
            if row_index == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F4F6F9")
                tc_pr.append(shd)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if row_index == 0
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in paragraph.runs:
                    set_run_font(run, size=8.3, bold=(row_index == 0))


def style_docx(path: Path) -> None:
    doc = Document(path)
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(22.45)
    section.right_margin = Mm(22.45)
    section.top_margin = Mm(24)
    section.bottom_margin = Mm(22)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)
    section.different_first_page_header_footer = True

    # narrative_proposal preset, exact core tokens
    normal = doc.styles["Normal"]
    set_style_font(normal, "微软雅黑", "Calibri", 11)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333
    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = doc.styles[name]
        set_style_font(style, "微软雅黑", "Calibri", size, bold=True, color=color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Paragraph", "List Bullet", "List Number"):
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        set_style_font(style, "微软雅黑", "Calibri", 11)
        style.paragraph_format.left_indent = Mm(9.525)  # 0.375 in
        style.paragraph_format.first_line_indent = Mm(-4.928)  # 0.194 in
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208
    if "Caption" in doc.styles:
        caption = doc.styles["Caption"]
        set_style_font(caption, "微软雅黑", "Calibri", 9.5, color="555555")
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(8)

    in_appendix = False
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 1" and paragraph.text.startswith("第十三章"):
            in_appendix = True
        if in_appendix and not paragraph.style.name.startswith("Heading"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif paragraph.style.name == "Normal":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if paragraph.style.name == "Normal":
            for run in paragraph.runs:
                if not paragraph._p.xpath(".//m:oMath"):
                    set_run_font(run)
    for table in doc.tables:
        set_table_geometry(table)

    # Quiet running header/footer.
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.text = "倾转旋翼机开环操纵稳定特性分析研究报告"
    for run in header.runs:
        set_run_font(run, size=9, color="777777")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(footer, " PAGE ", "1")
    for run in footer.runs:
        set_run_font(run, size=9, color="777777")

    # editorial_cover: title, rule, subtitle, report identity, date.
    anchor = doc.paragraphs[0]
    title_p = insert_before(anchor, TITLE)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(100)
    title_p.paragraph_format.space_after = Pt(28)
    for run in title_p.runs:
        set_run_font(run, size=28, bold=True, color="2E74B5")
    rule_p = insert_before(anchor, "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    rule_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in rule_p.runs:
        set_run_font(run, size=10, color="2E74B5")
    type_p = insert_before(anchor, "通用低阶部件级模型技术研究报告")
    type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in type_p.runs:
        set_run_font(run, size=15, color="1F4D78")
    date_p = insert_before(anchor, f"生成日期：{datetime.now():%Y年%m月%d日}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_break_cover = insert_before(anchor)
    page_break_cover.add_run().add_break(WD_BREAK.PAGE)
    toc_heading = insert_before(anchor, "目录")
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in toc_heading.runs:
        set_run_font(run, size=18, bold=True, color="2E74B5")
    toc = insert_before(anchor)
    add_field(toc, r' TOC \o "1-3" \h \z \u ', "更新域后显示目录")
    page_break_after_toc = insert_before(anchor)
    page_break_after_toc.add_run().add_break(WD_BREAK.PAGE)

    doc.core_properties.title = TITLE
    doc.core_properties.subject = "开环操纵稳定特性后处理与技术研究报告"
    doc.core_properties.author = "倾转旋翼机模型研究项目"
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    doc.save(path)


def _tab_run() -> OxmlElement:
    run = OxmlElement("w:r")
    run.append(OxmlElement("w:tab"))
    return run


def number_display_equations(path: Path) -> int:
    """Center editable OMML equations and add continuous right-side numbers."""
    doc = Document(path)
    equation_number = 0
    for paragraph in doc.paragraphs:
        math_para = paragraph._p.find(qn("m:oMathPara"))
        if math_para is None:
            continue
        equation = math_para.find(qn("m:oMath"))
        if equation is None:
            continue

        equation_number += 1
        math_para.remove(equation)
        index = paragraph._p.index(math_para)
        paragraph._p.remove(math_para)
        paragraph._p.insert(index, _tab_run())
        paragraph._p.insert(index + 1, equation)
        paragraph._p.insert(index + 2, _tab_run())

        number_run = OxmlElement("w:r")
        number_properties = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Times New Roman")
        fonts.set(qn("w:hAnsi"), "Times New Roman")
        fonts.set(qn("w:eastAsia"), "宋体")
        number_properties.append(fonts)
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), "21")
        number_properties.append(size)
        number_run.append(number_properties)
        number_text = OxmlElement("w:t")
        number_text.text = f"({equation_number})"
        number_run.append(number_text)
        paragraph._p.insert(index + 3, number_run)

        properties = paragraph._p.get_or_add_pPr()
        alignment = properties.find(qn("w:jc"))
        if alignment is not None:
            properties.remove(alignment)
        old_tabs = properties.find(qn("w:tabs"))
        if old_tabs is not None:
            properties.remove(old_tabs)
        tabs = OxmlElement("w:tabs")
        for value, position in (("center", "4680"), ("right", "9360")):
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), value)
            tab.set(qn("w:pos"), position)
            tabs.append(tab)
        properties.append(tabs)

    doc.save(path)
    return equation_number


def build_docx(source: str) -> None:
    command = [
        str(PANDOC),
        "--from=markdown+tex_math_dollars+tex_math_single_backslash",
        "--to=docx",
        f"--resource-path={REPORT}",
        "--output",
        str(DOCX_OUT),
    ]
    result = subprocess.run(
        command,
        input=source,
        text=True,
        encoding="utf-8",
        capture_output=True,
    )
    if result.returncode or result.stderr.strip():
        raise RuntimeError(result.stderr)
    style_docx(DOCX_OUT)
    display_count = number_display_equations(DOCX_OUT)
    expected_count = source.count("$$") // 2
    if expected_count != EXPECTED_DISPLAY_EQUATIONS:
        raise RuntimeError(
            "Source equation count changed: "
            f"expected={EXPECTED_DISPLAY_EQUATIONS}, actual={expected_count}"
        )
    if display_count != expected_count:
        raise RuntimeError(
            f"Display equation count mismatch: Word={display_count}, source={expected_count}"
        )


def docx_structure(path: Path) -> dict[str, int]:
    doc = Document(path)
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "figures": len(doc.inline_shapes),
        "equation_numbers": len(re.findall(r"<w:t>\(\d+\)</w:t>", xml)),
        "omml_total": xml.count("<m:oMath"),
        "raw_tex_markers": xml.count("$$"),
        "toc_fields": xml.count('TOC \\\\o') + xml.count('TOC \\o'),
    }


def main() -> None:
    missing = [name for name in REQUIRED_CSV if not (REPORT / name).exists()]
    if missing:
        raise SystemExit("Missing MATLAB evidence: " + ", ".join(missing))
    if not PANDOC.exists():
        raise SystemExit(f"Pandoc missing: {PANDOC}")
    (REPORT / "figures").mkdir(parents=True, exist_ok=True)
    source = normalize_word_math_source(build_markdown())
    validate_report_language(source)
    build_docx(source)
    structure = docx_structure(DOCX_OUT)
    log = [
        f"timestamp={datetime.now().isoformat()}",
        f"python={sys.version}",
        f"pandoc={PANDOC}",
        f"docx={DOCX_OUT}",
        *(f"{key}={value}" for key, value in structure.items()),
    ]
    print("\n".join(log))


if __name__ == "__main__":
    main()
