#!/usr/bin/env python3
"""Build an editable Word thesis from the frozen XeLaTeX submission source.

The conversion keeps prose, headings, tables, figures, citations and equations
editable. Pandoc creates native OMML equations; python-docx and OOXML patches
then apply an A4 Chinese academic-thesis style and deterministic numbering.
"""

from __future__ import annotations

import argparse
import re
import subprocess
import tempfile
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph


TITLE = "倾转旋翼机部件级飞行动力学建模、短舱动态状态扩展与可信度分析"
BODY_CHAPTERS = [
    "绪论",
    "坐标系、动力学理论与模型总体架构",
    "部件级非线性飞行动力学模型",
    "配平、数值线性化与可信度判据",
    "左右短舱动态状态扩展",
    "模型校核、外部数据关联与可信度方法",
    "模型校核与外部关联结果",
    "短舱动态状态对刚体响应的影响",
    "旋翼模型、参数敏感性与配平边界",
    "结论与展望",
]
APPENDICES = [
    "公式—代码—参数—测试追溯说明",
    "验证与外部数据矩阵",
    "可重复性与数据完整性",
    "状态、输入与复现实验说明",
]


def set_run_font(run, east_asia: str, latin: str, size: float, bold=None, italic=None):
    run.font.name = latin
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:cs"), latin)


def set_style_font(style, east_asia: str, latin: str, size: float, bold=None):
    style.font.name = latin
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:cs"), latin)


def insert_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph):
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, "宋体", "Times New Roman", 9)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths: list[int], table_width: int = 8784):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(table_width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_index, row in enumerate(table.rows):
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index == 0:
                shade_cell(cell, "EDEDED")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        "宋体",
                        "Times New Roman",
                        8.2 if len(widths) == 7 else 9.5,
                        bold=(row_index == 0),
                    )


def repair_longtable_conversion(table):
    """Remove Pandoc continuation headers and restore numeric-only cells.

    Pandoc converts LaTeX ``longtable`` continuation headers into a duplicate
    first row.  In the same conversion path, cells beginning with a bare number
    can be interpreted as list labels and lose that number.  The affected
    values below are copied verbatim from the frozen submission source.
    """
    if len(table.rows) >= 2:
        first = [cell.text.strip() for cell in table.rows[0].cells]
        second = [cell.text.strip() for cell in table.rows[1].cells]
        if first == second:
            table._tbl.remove(table.rows[1]._tr)

    header = [cell.text.strip() for cell in table.rows[0].cells]
    rows = None
    if header == [
        "工况",
        "短舱角/(°)",
        "速度/(m/s)",
        "俯仰姿态/(°)",
        "动态残差范数",
        "条件数",
        "最小控制余度",
    ]:
        rows = [
            ["B15/V20", "15", "20", "12.7200", "4.54×10^-10", "59.96", "0.2075"],
            ["B45/V35", "45", "35", "26.6719", "1.36×10^-9", "74.09", "0.1190"],
            ["B75/V80", "75", "80", "7.8512", "3.83×10^-9", "23.83", "0.3878"],
        ]
    elif header == [
        "工况",
        "输入",
        "最大俯仰角速度/(rad/s)",
        "最大滚转角速度/(rad/s)",
        "最大偏航角速度/(rad/s)",
        "动态配平偏离",
        "时间步峰值差",
    ]:
        rows = [
            ["15°/20 m/s", "对称2°", "0.00601", "0", "0", "0.7405", "1.5248%"],
            ["15°/20 m/s", "差动1°", "0.00402", "0.05490", "0.17885", "3.9688", "1.0728%"],
            ["45°/35 m/s", "对称2°", "0.01809", "0", "0", "0.7405", "1.3264%"],
            ["45°/35 m/s", "差动1°", "0.01034", "0.17773", "0.03587", "5.2308", "1.5145%"],
            ["75°/80 m/s", "对称2°", "0.00234", "0", "0", "0.7406", "0.4058%"],
            ["75°/80 m/s", "差动1°", "0.00101", "0.04884", "0.03017", "1.7792", "1.2111%"],
        ]
    elif header == ["方案", "主要变化", "九点可信/病态/失败", "论文角色", "不得解释为"]:
        rows = [
            ["原始通用基线", "概念参数", "7/0/2", "物理基线", "XV-15数据集"],
            ["公开参数覆盖", "有限公开参数", "5/2/2", "来源敏感性", "完整型号复现"],
            ["纵向几何优化", "三项布局量", "8/0/1", "配平能力研究", "参数辨识"],
            [
                "几何与等效控制联合优化",
                "布局量加升降舵等效效能",
                "8/0/1",
                "提供可信动态起点",
                "试验校准",
            ],
        ]

    if rows is not None:
        if len(table.rows) != len(rows) + 1:
            raise RuntimeError(f"unexpected converted table size for {header[0]}: {len(table.rows)}")
        for target_row, values in zip(table.rows[1:], rows):
            for cell, value in zip(target_row.cells, values):
                cell.text = value


def number_display_equations(doc):
    """Add editable, chapter-scoped numbers to native OMML display equations."""
    chapter_number = None
    appendix_letter = None
    equation_count = 0
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 1":
            text = paragraph.text.strip()
            body_match = re.match(r"第(\d+)章", text)
            app_match = re.match(r"附录([A-Z])", text)
            if body_match:
                chapter_number = body_match.group(1)
                appendix_letter = None
                equation_count = 0
            elif app_match:
                chapter_number = None
                appendix_letter = app_match.group(1)
                equation_count = 0
            continue
        if not paragraph._p.xpath(".//m:oMathPara"):
            continue
        if chapter_number is None and appendix_letter is None:
            continue
        equation_count += 1
        prefix = chapter_number if chapter_number is not None else appendix_letter
        paragraph.paragraph_format.tab_stops.add_tab_stop(Mm(155), WD_TAB_ALIGNMENT.RIGHT)
        run = paragraph.add_run(f"\t({prefix}.{equation_count})")
        set_run_font(run, "宋体", "Times New Roman", 10.5)


def parse_bib(path: Path) -> OrderedDict[str, dict[str, str]]:
    text = path.read_text(encoding="utf-8")
    entries: OrderedDict[str, dict[str, str]] = OrderedDict()
    pos = 0
    while True:
        match = re.search(r"@(\w+)\s*\{\s*([^,\s]+)\s*,", text[pos:], re.S)
        if not match:
            break
        entry_type, key = match.group(1).lower(), match.group(2)
        start = pos + match.end()
        depth = 1
        index = start
        while index < len(text) and depth:
            if text[index] == "{":
                depth += 1
            elif text[index] == "}":
                depth -= 1
            index += 1
        body = text[start : index - 1]
        fields: dict[str, str] = {"entrytype": entry_type, "key": key}
        cursor = 0
        while cursor < len(body):
            field_match = re.search(r"(\w+)\s*=\s*\{", body[cursor:], re.S)
            if not field_match:
                break
            name = field_match.group(1).lower()
            value_start = cursor + field_match.end()
            value_depth = 1
            value_end = value_start
            while value_end < len(body) and value_depth:
                if body[value_end] == "{":
                    value_depth += 1
                elif body[value_end] == "}":
                    value_depth -= 1
                value_end += 1
            value = body[value_start : value_end - 1]
            value = value.replace(r"\&", "&").replace("--", "–").strip()
            fields[name] = value
            cursor = value_end
        entries[key] = fields
        pos = index
    return entries


def author_text(value: str) -> str:
    if not value:
        return "UNKNOWN"
    authors = []
    for raw in value.split(" and "):
        raw = raw.strip().strip("{}")
        if "," in raw:
            family, given = [part.strip() for part in raw.split(",", 1)]
            initials = "".join(part[0].upper() for part in re.findall(r"[A-Za-z]+", given))
            authors.append(f"{family.upper()} {initials}".strip())
        else:
            authors.append(raw.upper())
    return ", ".join(authors)


def format_reference(entry: dict[str, str]) -> str:
    kind = entry.get("entrytype", "")
    author = author_text(entry.get("author", ""))
    title = entry.get("title", "").replace("{", "").replace("}", "")
    year = entry.get("year", "")
    if kind == "article":
        journal = entry.get("journal", "")
        volume = entry.get("volume", "")
        number = entry.get("number", "")
        pages = entry.get("pages", "")
        volume_issue = volume + (f"({number})" if number else "")
        tail = f"{journal}, {year}, {volume_issue}"
        if pages:
            tail += f": {pages}"
        doi = entry.get("doi", "")
        if doi:
            tail += f". DOI: {doi}"
        return f"{author}. {title}[J]. {tail}."
    if kind in {"book", "incollection"}:
        mark = "M"
        edition = entry.get("edition", "")
        address = entry.get("address", "")
        publisher = entry.get("publisher", "")
        extra = f"{edition}nd ed. " if edition == "2" else (f"{edition}rd ed. " if edition == "3" else (f"{edition}th ed. " if edition else ""))
        if kind == "incollection":
            booktitle = entry.get("booktitle", "")
            extra = f"In: {booktitle}. "
        place_pub = ": ".join(part for part in (address, publisher) if part)
        return f"{author}. {title}[{mark}]. {extra}{place_pub}, {year}."
    if kind == "phdthesis":
        return f"{author}. {title}[D]. {entry.get('school', '')}, {year}."
    if kind == "techreport":
        number = entry.get("number", "")
        institution = entry.get("institution", "")
        return f"{author}. {title}[R]. {institution}, {number}, {year}."
    if kind == "standard":
        number = entry.get("number", "")
        institution = entry.get("institution", "")
        tail = ", ".join(part for part in (number, institution, year) if part)
        return f"{author}. {title}[S]. {tail}."
    return f"{author}. {title}. {year}."


def compact_numbers(numbers: list[int]) -> str:
    values = sorted(dict.fromkeys(numbers))
    groups = []
    start = previous = values[0]
    for number in values[1:]:
        if number == previous + 1:
            previous = number
            continue
        groups.append(str(start) if start == previous else f"{start}–{previous}")
        start = previous = number
    groups.append(str(start) if start == previous else f"{start}–{previous}")
    return "[" + ",".join(groups) + "]"


def prepare_tex(source: str, entries: OrderedDict[str, dict[str, str]]):
    cite_keys: list[str] = []
    for match in re.finditer(r"\\cite\{([^}]+)\}", source):
        for key in match.group(1).split(","):
            key = key.strip()
            if key and key not in cite_keys:
                cite_keys.append(key)
    for key in entries:
        if key not in cite_keys:
            cite_keys.append(key)
    citation_numbers = {key: index + 1 for index, key in enumerate(cite_keys)}

    def cite_replace(match):
        numbers = [
            citation_numbers[key.strip()]
            for key in match.group(1).split(",")
            if key.strip() in citation_numbers
        ]
        return compact_numbers(numbers)

    cleaned = re.sub(r"\\cite\{([^}]+)\}", cite_replace, source)
    cleaned = re.sub(
        r"\{\\rm\s+([^{}]+)\}",
        lambda match: r"\mathrm{" + match.group(1).strip() + "}",
        cleaned,
    )
    cleaned = re.sub(
        r"\\rm\s+([A-Za-z]+)",
        lambda match: r"\mathrm{" + match.group(1) + "}",
        cleaned,
    )
    cleaned = re.sub(r"\\emph\{图\s*[^{}]*\}\s*", "", cleaned)
    cleaned = re.sub(
        r"\\printbibliography(?:\[[^\]]*\])?",
        r"\\chapter*{参考文献}",
        cleaned,
    )
    return cleaned, cite_keys


def style_document(doc: Document, cite_keys: list[str], entries: OrderedDict[str, dict[str, str]]):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(30)
    section.right_margin = Mm(25)
    section.top_margin = Mm(28)
    section.bottom_margin = Mm(25)
    section.header_distance = Mm(12)
    section.footer_distance = Mm(12)
    section.different_first_page_header_footer = True

    styles = doc.styles
    for name in ("Normal", "Body Text", "First Paragraph"):
        if name not in styles:
            continue
        style = styles[name]
        set_style_font(style, "宋体", "Times New Roman", 12)
        fmt = style.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.first_line_indent = Pt(24)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1.5

    heading_tokens = {
        "Heading 1": ("黑体", "Times New Roman", 16, 24, 18, WD_ALIGN_PARAGRAPH.CENTER),
        "Heading 2": ("黑体", "Times New Roman", 14, 12, 6, WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 3": ("黑体", "Times New Roman", 12, 8, 4, WD_ALIGN_PARAGRAPH.LEFT),
    }
    for name, (east, latin, size, before, after, align) in heading_tokens.items():
        style = styles[name]
        set_style_font(style, east, latin, size, bold=True)
        style.font.color.rgb = RGBColor(0, 0, 0)
        fmt = style.paragraph_format
        fmt.alignment = align
        fmt.first_line_indent = Pt(0)
        fmt.space_before = Pt(before)
        fmt.space_after = Pt(after)
        fmt.keep_with_next = True
        if name == "Heading 1":
            fmt.page_break_before = True

    for caption_style in ("Caption", "Image Caption", "Table Caption"):
        if caption_style in styles:
            style = styles[caption_style]
            set_style_font(style, "宋体", "Times New Roman", 10.5)
            fmt = style.paragraph_format
            fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.first_line_indent = Pt(0)
            fmt.space_before = Pt(4)
            fmt.space_after = Pt(6)
            fmt.line_spacing = 1.0

    # Cover uses an editorial-cover pattern with an academic-title override.
    cover = doc.paragraphs[:3]
    if len(cover) >= 3:
        cover[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover[0].paragraph_format.space_before = Pt(80)
        cover[0].paragraph_format.space_after = Pt(42)
        cover[0].paragraph_format.first_line_indent = Pt(0)
        for run in cover[0].runs:
            set_run_font(run, "黑体", "Times New Roman", 22, bold=True)
        cover[1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover[1].paragraph_format.space_after = Pt(150)
        cover[1].paragraph_format.first_line_indent = Pt(0)
        for run in cover[1].runs:
            set_run_font(run, "宋体", "Times New Roman", 16)
        cover[2].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover[2].paragraph_format.first_line_indent = Pt(0)
        cover[2].paragraph_format.page_break_after = True
        for run in cover[2].runs:
            set_run_font(run, "宋体", "Times New Roman", 12)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = TITLE
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hp.runs:
        set_run_font(run, "宋体", "Times New Roman", 9)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    add_page_field(fp)

    # Number chapters and sections while retaining true Word heading styles.
    current_chapter: int | None = None
    current_appendix: str | None = None
    section_counter = 0
    body_lookup = {name: i + 1 for i, name in enumerate(BODY_CHAPTERS)}
    appendix_lookup = {name: chr(65 + i) for i, name in enumerate(APPENDICES)}
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 1":
            raw = paragraph.text.strip()
            current_chapter = body_lookup.get(raw)
            current_appendix = appendix_lookup.get(raw)
            section_counter = 0
            if current_chapter is not None:
                paragraph.text = f"第{current_chapter}章　{raw}"
            elif current_appendix is not None:
                paragraph.text = f"附录{current_appendix}　{raw}"
            paragraph.paragraph_format.page_break_before = True
            for run in paragraph.runs:
                set_run_font(run, "黑体", "Times New Roman", 16, bold=True)
        elif paragraph.style.name == "Heading 2":
            raw = paragraph.text.strip()
            section_counter += 1
            if current_chapter is not None:
                paragraph.text = f"{current_chapter}.{section_counter}　{raw}"
            elif current_appendix is not None:
                paragraph.text = f"{current_appendix}.{section_counter}　{raw}"
            for run in paragraph.runs:
                set_run_font(run, "黑体", "Times New Roman", 14, bold=True)

    number_display_equations(doc)

    # Number figure and table captions by chapter. Pre-main captions use
    # simple sequential numbers, mirroring the XeLaTeX front matter.
    chapter_number = None
    appendix_letter = None
    figure_count = 0
    table_count = 0
    front_figure = 0
    front_table = 0
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 1":
            text = paragraph.text.strip()
            body_match = re.match(r"第(\d+)章", text)
            app_match = re.match(r"附录([A-Z])", text)
            if body_match:
                chapter_number = int(body_match.group(1))
                appendix_letter = None
                figure_count = table_count = 0
            elif app_match:
                chapter_number = None
                appendix_letter = app_match.group(1)
                figure_count = table_count = 0
        elif paragraph.style.name == "Image Caption":
            label = re.sub(r"^图\s*[\w.]+\s*", "", paragraph.text.strip())
            if chapter_number is not None:
                figure_count += 1
                prefix = f"图 {chapter_number}.{figure_count}　"
            elif appendix_letter is not None:
                figure_count += 1
                prefix = f"图 {appendix_letter}.{figure_count}　"
            else:
                front_figure += 1
                prefix = f"图 {front_figure}　"
            paragraph.text = prefix + label
        elif paragraph.style.name == "Table Caption":
            label = re.sub(r"^表\s*[\w.]+\s*", "", paragraph.text.strip())
            if chapter_number is not None:
                table_count += 1
                prefix = f"表 {chapter_number}.{table_count}　"
            elif appendix_letter is not None:
                table_count += 1
                prefix = f"表 {appendix_letter}.{table_count}　"
            else:
                front_table += 1
                prefix = f"表 {front_table}　"
            paragraph.text = prefix + label

    # Use the editable figure captions as meaningful image alternative text.
    # Pandoc names extracted assets generically (image1.png, image2.png, ...),
    # so filename-derived descriptions would not help screen-reader users.
    figure_captions = [
        paragraph.text.strip()
        for paragraph in doc.paragraphs
        if paragraph.style.name == "Image Caption"
    ]
    if len(figure_captions) != len(doc.inline_shapes):
        raise RuntimeError(
            "Figure-caption count does not match inline-image count: "
            f"{len(figure_captions)} captions vs {len(doc.inline_shapes)} images"
        )
    for shape, caption in zip(doc.inline_shapes, figure_captions):
        shape._inline.docPr.set("title", caption)
        shape._inline.docPr.set("descr", caption)

    # Insert an editable Word TOC field and static figure/table directories
    # before the Chinese abstract, after caption numbers are materialized.
    declaration_text = next(
        (p for p in doc.paragraphs if p.text.strip().startswith("本页为占位页")),
        None,
    )
    if declaration_text is not None:
        toc_title = insert_after(declaration_text, "目录", "Heading 1")
        toc_title.paragraph_format.page_break_before = True
        toc_placeholder = insert_after(toc_title, "[[TOC]]", "Normal")
        toc_placeholder.paragraph_format.first_line_indent = Pt(0)

        figure_title = insert_after(toc_placeholder, "图目录", "Heading 1")
        figure_lines = []
        table_lines = []
        for paragraph in doc.paragraphs:
            if paragraph.style.name == "Image Caption":
                figure_lines.append(paragraph.text.strip())
            elif paragraph.style.name == "Table Caption":
                table_lines.append(paragraph.text.strip())
        cursor = figure_title
        for text in figure_lines:
            cursor = insert_after(cursor, text, "Normal")
            cursor.paragraph_format.first_line_indent = Pt(0)
            cursor.paragraph_format.line_spacing = 1.15
        table_title = insert_after(cursor, "表目录", "Heading 1")
        cursor = table_title
        for text in table_lines:
            cursor = insert_after(cursor, text, "Normal")
            cursor.paragraph_format.first_line_indent = Pt(0)
            cursor.paragraph_format.line_spacing = 1.15

    # Center figures/equations and keep figures with captions.
    for paragraph in doc.paragraphs:
        has_drawing = bool(paragraph._p.xpath(".//w:drawing"))
        has_display_math = bool(paragraph._p.xpath(".//m:oMathPara"))
        if has_drawing or has_display_math:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
        if has_drawing:
            paragraph.paragraph_format.keep_with_next = True

    # English abstract must not use Chinese first-line indentation.
    in_english_abstract = False
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 1":
            text = paragraph.text.strip()
            in_english_abstract = text == "Abstract"
            if in_english_abstract:
                continue
            if text not in {"Abstract"}:
                in_english_abstract = False
        if in_english_abstract and paragraph.style.name in {"Body Text", "First Paragraph", "Normal"}:
            paragraph.paragraph_format.first_line_indent = Pt(0)
            for run in paragraph.runs:
                set_run_font(run, "宋体", "Times New Roman", 11)

    # Add the 40 editable numeric references at the bibliography heading.
    reference_heading = next(
        (p for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.strip() == "参考文献"),
        None,
    )
    if reference_heading is not None:
        cursor = reference_heading
        for index, key in enumerate(cite_keys, start=1):
            entry = entries.get(key)
            if not entry:
                continue
            cursor = insert_after(cursor, f"[{index}] {format_reference(entry)}", "Normal")
            cursor.paragraph_format.first_line_indent = Pt(0)
            cursor.paragraph_format.left_indent = Pt(20)
            cursor.paragraph_format.first_line_indent = Pt(-20)
            cursor.paragraph_format.space_after = Pt(3)
            cursor.paragraph_format.line_spacing = 1.15
            for run in cursor.runs:
                set_run_font(run, "宋体", "Times New Roman", 10)

    # Explicit fixed table geometry (A4 usable width = 155 mm = 8784 DXA).
    for table in doc.tables:
        repair_longtable_conversion(table)
        columns = len(table.columns)
        if columns == 2:
            widths = [2100, 6684]
        elif columns == 3:
            widths = [1900, 5000, 1884]
        elif columns == 7:
            widths = [1150, 900, 1100, 1100, 1100, 1500, 1934]
        elif columns == 4:
            widths = [1500, 2200, 2400, 2684]
        else:
            base = 8784 // columns
            widths = [base] * columns
            widths[-1] += 8784 - sum(widths)
        set_table_geometry(table, widths)

    doc.core_properties.title = TITLE
    doc.core_properties.subject = "可编辑Word提交候选稿"
    doc.core_properties.author = "匿名送审稿"
    doc.core_properties.comments = "由最终XeLaTeX提交稿转换；正文、公式、图表和参考文献均可编辑。"
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--tex", type=Path, required=True)
    parser.add_argument("--bib", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--pandoc", type=Path, required=True)
    parser.add_argument("--toc-script", type=Path, required=True)
    args = parser.parse_args()

    args.output.parent.mkdir(parents=True, exist_ok=True)
    entries = parse_bib(args.bib)
    source = args.tex.read_text(encoding="utf-8")
    cleaned, cite_keys = prepare_tex(source, entries)

    with tempfile.TemporaryDirectory(prefix="tiltrotor_word_") as tmp:
        tmp_dir = Path(tmp)
        clean_tex = tmp_dir / "thesis_word_source.tex"
        baseline = tmp_dir / "pandoc.docx"
        styled = tmp_dir / "styled.docx"
        clean_tex.write_text(cleaned, encoding="utf-8")
        command = [
            str(args.pandoc),
            str(clean_tex),
            "--from=latex",
            "--to=docx",
            f"--resource-path={args.tex.parent}",
            "-o",
            str(baseline),
        ]
        result = subprocess.run(command, text=True, capture_output=True, check=False)
        if result.stdout:
            print(result.stdout)
        if result.stderr:
            print(result.stderr)
        if result.returncode:
            raise SystemExit(result.returncode)

        doc = Document(baseline)
        style_document(doc, cite_keys, entries)
        doc.save(styled)

        # The current Python runtime is used by the caller; call the helper as
        # a module script with sys.executable for deterministic dependencies.
        import sys

        result = subprocess.run(
            [sys.executable, str(args.toc_script), str(styled), "--out", str(args.output)],
            text=True,
            capture_output=True,
            check=False,
        )
        if result.stdout:
            print(result.stdout)
        if result.stderr:
            print(result.stderr)
        if result.returncode:
            raise SystemExit(result.returncode)

    final_doc = Document(args.output)
    xml = final_doc._element.xml
    print(
        {
            "paragraphs": len(final_doc.paragraphs),
            "tables": len(final_doc.tables),
            "figures": len(final_doc.inline_shapes),
            "references": len(cite_keys),
            "omml_display": xml.count("<m:oMathPara"),
            "omml_total": xml.count("<m:oMath"),
            "output": str(args.output),
        }
    )


if __name__ == "__main__":
    main()
